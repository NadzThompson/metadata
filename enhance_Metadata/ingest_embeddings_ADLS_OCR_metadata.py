# Databricks notebook source
# ingest_embeddings_ADLS_OCR_metadata.py
#
# NOVA RAG Pipeline -- Ingestion, OCR & Embedding (Databricks)
#
# Pipeline Flow:
#   File -> OCR/Parse -> Chunk -> Enrich (NOVA 3-Rule) -> Embed -> Elasticsearch + PGVector
#
# Steps:
#   1. Initialize OAuth token manager
#   2. Fetch files from ADLS (Azure Data Lake Storage)
#   3. Connect to Elasticsearch
#   4. Create/verify Elasticsearch index with NOVA mapping
#   5. Check which files already exist in the index
#   6. Process and chunk documents (PDF/DOCX/HTML/TXT with OCR)
#   6b. Apply Rule 1 semantic headers and assign chunk_ids
#   7. Generate embeddings
#   8. Ingest into Elasticsearch (with NOVA INDEX_FIELDS at top level)
#   9. PGVector dual-store upsert
#
# TABLE OF CONTENTS
# -----------------
#  Cell  1: Header docstring (this cell)
#  Cell  2: Imports
#  Cell  3: OCR Processor Import
#  Cell  4: Logging
#  Cell  5: Configuration & Environment Variables
#  Cell  6: NOVA Three-Rule Constants
#  Cell  7: NOVA Dataclasses
#  Cell  8: ProxyManager
#  Cell  9: NOVA Structural Metadata Helpers
#  Cell 10: NOVA Semantic Header & Prompt Rendering
#  Cell 11: NOVA Metadata Enrichment
#  Cell 12: Regulatory Scraped JSON Parser
#  Cell 13: Elasticsearch Connection & Index Management
#  Cell 14: PGVector Dual-Store Support
#  Cell 15: ADLS File Operations
#  Cell 16: Check File Existence
#  Cell 17: HTML Processing Helpers
#  Cell 18: PDF Processing with OCR
#  Cell 19: DOCX Processing with OCR
#  Cell 20: Text Chunking
#  Cell 21: HTML Processing with OCR
#  Cell 22: Main Pipeline - load_and_split_documents()
#  Cell 23: Main Entry Point
#  Cell 24: Azure Functions Entry Point
#  Cell 25: __main__ block

# COMMAND ----------

import json
import os
import logging
import tempfile
import hashlib
import datetime
import re
import io
import base64
import traceback
from pathlib import Path
from typing import List, Dict, Optional, Tuple
from dataclasses import dataclass, field as dc_field

from openai import OpenAI
from elasticsearch import Elasticsearch, helpers

from azure.storage.blob import BlobServiceClient
from azure.identity import ClientSecretCredential

from langchain.text_splitter import RecursiveCharacterTextSplitter, MarkdownHeaderTextSplitter
from langchain.schema import Document
from langchain_community.document_loaders import TextLoader

from bs4 import BeautifulSoup, NavigableString, Tag

import fitz  # PyMuPDF

try:
    import azure.functions as func
    AZURE_FUNCTIONS_AVAILABLE = True
except ImportError:
    func = None
    AZURE_FUNCTIONS_AVAILABLE = False

try:
    import httpx
except ImportError:
    httpx = None

# COMMAND ----------

try:
    from ocr_processor import (
        DocumentOCRProcessor,
        OAuthTokenManager,
        is_pdf_text_only,
        is_pdf_scanned,
        PDF_SUPPORT,
        DOCX_SUPPORT,
    )
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False
    OAuthTokenManager = None
    DocumentOCRProcessor = None

# COMMAND ----------

def log_and_print(message, level="info"):
    """Print message with UTC timestamp and log it.

    Args:
        message: The message string to log.
        level: Logging level -- "info", "warning", or "error".
    """
    timestamp = datetime.datetime.utcnow().strftime("%Y-%m-%d %H:%M:%S")
    print(f"[{timestamp}] {message}")
    if level == "info":
        logging.info(message)
    elif level == "warning":
        logging.warning(message)
    elif level == "error":
        logging.error(message)


logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s - %(name)s - %(levelname)s - %(message)s",
)

logger = logging.getLogger("ingest_adls_ocr")

logging.getLogger("azure").setLevel(logging.WARNING)
logging.getLogger("azure.core.pipeline.policies.http_logging_policy").setLevel(logging.WARNING)
logging.getLogger("azure.storage.blob").setLevel(logging.WARNING)
logging.getLogger("azure.identity").setLevel(logging.WARNING)
logging.getLogger("elasticsearch").setLevel(logging.WARNING)
logging.getLogger("elastic_transport").setLevel(logging.WARNING)
logging.getLogger("urllib3").setLevel(logging.WARNING)
logging.getLogger("requests").setLevel(logging.WARNING)

# COMMAND ----------

# === Configuration & Environment Variables ===

ELASTICSEARCH_HOST = os.environ.get("ELASTICSEARCH_HOST", "")
ELASTICSEARCH_INDEX = os.environ.get("ELASTICSEARCH_INDEX", "ctknowledgebasemetadata2")
ELASTICSEARCH_CA_FINGERPRINT = os.environ.get("ELASTICSEARCH_CA_FINGERPRINT", "")

AZURE_TENANT_ID = os.environ.get("AZURE_TENANT_ID", "")
AZURE_CLIENT_ID = os.environ.get("AZURE_CLIENT_ID", "")
AZURE_CLIENT_SECRET = os.environ.get("AZURE_CLIENT_SECRET", "")
ADLS_URL = os.environ.get("ADLS_URL", "")

ELASTICSEARCH_CLOUD_ID = os.environ.get("ELASTICSEARCH_CLOUD_ID", "")
ELASTICSEARCH_API_KEY = os.environ.get("ELASTICSEARCH_API_KEY", "")
ELASTICSEARCH_BASIC_USER = os.environ.get("ELASTICSEARCH_BASIC_USER", "elastic")
ELASTICSEARCH_BASIC_PASS = os.environ.get("ELASTICSEARCH_BASIC_PASS", "")

AZURE_STORAGE_ACCOUNT_URL = os.environ.get("AZURE_STORAGE_ACCOUNT_URL", ADLS_URL)
AZURE_STORAGE_CONTAINER = os.environ.get("AZURE_STORAGE_CONTAINER", "lwf0-ingress")
AZURE_STORAGE_FOLDER = os.environ.get("AZURE_STORAGE_FOLDER", "lwf0-ingress/ctknowledgebase/testmetadata")

EMBEDDING_MODEL = os.environ.get("EMBEDDING_MODEL", "text-embedding-3-large")
EMBEDDING_DIMS = 1024  # text-embedding-3-large default
CHUNK_SIZE = int(os.environ.get("CHUNK_SIZE", "1500"))
CHUNK_OVERLAP = int(os.environ.get("CHUNK_OVERLAP", "100"))

INDEX_NAME = ELASTICSEARCH_INDEX or "ctknowledgebasemetadata2"
container_name = AZURE_STORAGE_CONTAINER or "lwf0-ingress"
storage_folder = AZURE_STORAGE_FOLDER or "lwf0-ingress/ctknowledgebase/testmetadata"

VISION_MODEL = os.environ.get("VISION_MODEL", "gpt-4o")

# PGVector env vars (optional dual-store)
PG_HOST = os.environ.get("PG_HOST", "")
PG_PORT = int(os.environ.get("PG_PORT", "5432"))
PG_DB = os.environ.get("PG_DB", "nova")
PG_USER = os.environ.get("PG_USER", "nova")
PG_PASSWORD = os.environ.get("PG_PASSWORD", "")
PG_TABLE = os.environ.get("PG_TABLE", "nova_chunks")

# COMMAND ----------

# === NOVA Three-Rule Constants ===

# ADLS layout prefixes
BRONZE_EXTERNAL_PREFIX = "bronze/external/"
BRONZE_INTERNAL_PREFIX = "bronze/internal/"
SILVER_METADATA_PREFIX = "silver/metadata/"
GOLD_PREFIX = "gold/"

# Regulatory bodies known to the pipeline
KNOWN_REGULATORS = [
    {"acronym": "OSFI", "name": "Office of the Superintendent of Financial Institutions", "jurisdiction": "CA"},
    {"acronym": "PRA", "name": "Prudential Regulation Authority", "jurisdiction": "GB"},
    {"acronym": "OCC", "name": "Office of the Comptroller of the Currency", "jurisdiction": "US"},
    {"acronym": "FDIC", "name": "Federal Deposit Insurance Corporation", "jurisdiction": "US"},
    {"acronym": "FRB", "name": "Federal Reserve Board", "jurisdiction": "US"},
    {"acronym": "BCBS", "name": "Basel Committee on Banking Supervision", "jurisdiction": "INTL"},
    {"acronym": "ECB", "name": "European Central Bank", "jurisdiction": "EU"},
    {"acronym": "EBA", "name": "European Banking Authority", "jurisdiction": "EU"},
    {"acronym": "APRA", "name": "Australian Prudential Regulation Authority", "jurisdiction": "AU"},
    {"acronym": "MAS", "name": "Monetary Authority of Singapore", "jurisdiction": "SG"},
]

# Rule 1 -- EMBEDDED_FIELDS: prepended as a semantic header into the chunk text
#            before embedding so the vector captures these signals.
EMBEDDED_FIELDS = [
    "regulator_acronym",
    "guideline_number",
    "document_class",
    "section_path",
    "section_number",
    "normative_weight",
]

# Rule 2 -- INDEX_FIELDS: stored as ES / PG columns for filter and faceted search
#            but NOT embedded into the vector.
INDEX_FIELDS = [
    "source_type", "short_title", "document_class", "section_path",
    "citation_anchor", "status", "effective_date_start", "effective_date_end",
    "jurisdiction", "authority_class", "authority_level", "nova_tier",
    "regulator", "regulator_acronym", "guideline_number", "version_id",
    "version_label", "current_version_flag", "sector",
    "doc_family_id", "business_owner", "business_line", "audience",
    "approval_status", "confidentiality", "structural_level",
    "section_number", "depth", "normative_weight", "paragraph_role",
    "is_appendix", "cross_references",
    "contains_definition", "contains_formula", "contains_requirement",
    "contains_deadline", "contains_assignment", "contains_parameter",
    "bm25_text",
]

# Rule 3 -- PROMPT_INJECTED_FIELDS: rendered into the LLM context at inference time.
PROMPT_INJECTED_FIELDS_REGULATORY = [
    "title", "citation_anchor",
    "regulator", "regulator_acronym", "guideline_number",
    "short_title", "effective_date_start", "status",
    "version_id", "version_label", "current_version_flag",
    "authority_class", "nova_tier", "jurisdiction",
    "section_path", "section_number", "normative_weight",
    "paragraph_role", "cross_references",
]
PROMPT_INJECTED_FIELDS_INTERNAL = [
    "title", "citation_anchor",
    "business_owner", "business_line", "audience",
    "approval_status", "confidentiality",
    "version_id", "current_version_flag",
    "section_path", "section_number", "normative_weight",
    "paragraph_role",
]

# COMMAND ----------

# === NOVA Dataclasses ===


@dataclass
class CanonicalDocument:
    """Document-level metadata container (one per ingested file)."""
    doc_id: str = ""
    source_type: str = ""           # regulatory / internal / auto
    title: str = ""
    short_title: str = ""
    document_class: str = ""        # guideline / advisory / rule / policy / procedure / standard
    regulator: str = ""
    regulator_acronym: str = ""
    guideline_number: str = ""
    jurisdiction: str = ""
    authority_class: str = ""       # prudential / conduct / market / other
    authority_level: str = ""       # primary_legislation / regulation / guidance / internal_policy
    nova_tier: str = ""             # tier_1_binding / tier_2_expected / tier_3_guidance / internal
    status: str = "active"
    effective_date_start: str = ""
    effective_date_end: str = ""
    version_id: str = ""
    version_label: str = ""
    current_version_flag: bool = True
    sector: str = ""
    doc_family_id: str = ""
    business_owner: str = ""
    business_line: str = ""
    audience: str = ""
    approval_status: str = ""
    confidentiality: str = ""
    source_path: str = ""


@dataclass
class CanonicalUnit:
    """Chunk-level metadata container (one per text chunk)."""
    chunk_id: str = ""
    doc_id: str = ""
    chunk_index: int = 0
    total_chunks: int = 0
    section_path: str = ""
    section_number: str = ""
    citation_anchor: str = ""
    structural_level: str = ""      # chapter / section / subsection / paragraph / appendix
    depth: int = 0
    normative_weight: str = ""      # mandatory / advisory / permissive / informational
    paragraph_role: str = ""        # definition / requirement / procedure_step / example / exception / narrative
    is_appendix: bool = False
    cross_references: List[str] = dc_field(default_factory=list)
    contains_definition: bool = False
    contains_formula: bool = False
    contains_requirement: bool = False
    contains_deadline: bool = False
    contains_assignment: bool = False
    contains_parameter: bool = False
    bm25_text: str = ""
    page_number: int = 0
    heading_path: str = ""
    content_type: str = "text"

# COMMAND ----------

# === ProxyManager ===


class ProxyManager:
    """Manage proxy settings for ADLS and Elasticsearch connections."""

    BYPASS_ENABLED = False
    HTTPS_PROXY = None
    HTTP_PROXY = None
    NO_PROXY = None

    @staticmethod
    def clear_proxies_for_initialization():
        """Temporarily remove proxy settings from os.environ for ADLS connections."""
        proxy_keys = ["HTTP_PROXY", "HTTPS_PROXY", "http_proxy", "https_proxy"]
        for key in proxy_keys:
            if key in os.environ:
                del os.environ[key]

    @staticmethod
    def restore_proxies_for_elasticsearch():
        """Restore proxy settings after ADLS connection is established."""
        pass

    @staticmethod
    def set_no_proxy():
        """Set NO_PROXY to bypass all proxy settings."""
        os.environ["NO_PROXY"] = "*"

# COMMAND ----------

# === NOVA Structural Metadata Helpers ===


def _classify_normative_weight(text: str) -> str:
    """Classify the normative weight of a text chunk by scanning for modal verbs.

    Args:
        text: The chunk text to classify.

    Returns:
        One of "mandatory", "advisory", "permissive", or "informational".
    """
    lower = text.lower()
    if re.search(r'\b(shall|must|is required to|are required to)\b', lower):
        return "mandatory"
    if re.search(r'\b(should|expected to|is expected to|are expected to)\b', lower):
        return "advisory"
    if re.search(r'\b(may|can|is permitted to|are permitted to)\b', lower):
        return "permissive"
    return "informational"


def _classify_paragraph_role(text: str, heading: str = "") -> str:
    """Classify the role of a paragraph within a document.

    Args:
        text: The paragraph text to classify.
        heading: Optional heading text for additional context.

    Returns:
        One of "definition", "requirement", "procedure_step", "example",
        "exception", or "narrative".
    """
    lower = text.lower()
    heading_lower = heading.lower() if heading else ""

    if re.search(r'\b(means|refers to|is defined as|definition)\b', lower):
        return "definition"
    if "example" in heading_lower or re.search(r'\b(for example|e\.g\.|for instance|illustration)\b', lower):
        return "example"
    if re.search(r'\b(exception|notwithstanding|except where|unless)\b', lower):
        return "exception"
    if re.search(r'(step\s*\d|procedure|process\s*:)', lower) or re.search(r'^\s*\d+\.\s', text):
        return "procedure_step"
    if re.search(r'\b(shall|must|is required to|should|are required to)\b', lower):
        return "requirement"
    return "narrative"


def _extract_cross_references(text: str) -> List[str]:
    """Extract cross-references from text (e.g. 'see Section X', 'B-20', 'IFRS 9').

    Args:
        text: Source text to scan for references.

    Returns:
        Deduplicated list of cross-reference strings.
    """
    refs = []
    for m in re.finditer(r'(?:see|refer to|as per|per|under)\s+(Section|Chapter|Appendix|Annex|Part)\s+[\w\.\-]+', text, re.IGNORECASE):
        refs.append(m.group(0).strip())
    for m in re.finditer(r'\b([A-Z]{1,5}[\-\s]?\d{1,4}(?:\.\d+)*)\b', text):
        candidate = m.group(1).strip()
        if len(candidate) >= 3 and not candidate.replace("-", "").replace(".", "").isdigit():
            refs.append(candidate)
    return list(dict.fromkeys(refs))


def _extract_section_number(heading: str) -> str:
    """Extract a section number like '3.2.1' from a heading string.

    Args:
        heading: Heading text potentially starting with a dotted number.

    Returns:
        The extracted section number string, or "" if none found.
    """
    m = re.match(r'^\s*((?:\d+\.)*\d+)', heading)
    return m.group(1) if m else ""


def _infer_structural_level(depth: int, heading: str = "") -> str:
    """Infer structural level from depth and heading text.

    Args:
        depth: Nesting depth (0 = top-level).
        heading: Optional heading text for keyword detection.

    Returns:
        One of "chapter", "section", "subsection", "paragraph", or "appendix".
    """
    heading_lower = heading.lower() if heading else ""
    if "appendix" in heading_lower or "annex" in heading_lower:
        return "appendix"
    if depth == 0 or "chapter" in heading_lower:
        return "chapter"
    if depth == 1 or re.match(r'^\d+\s', heading):
        return "section"
    if depth == 2:
        return "subsection"
    return "paragraph"


def _compute_content_flags(text: str) -> Dict[str, bool]:
    """Compute boolean content flags for a text chunk.

    Args:
        text: Chunk text to scan for semantic markers.

    Returns:
        Dict mapping flag names (e.g. "contains_definition") to booleans.
    """
    lower = text.lower()
    return {
        "contains_definition": bool(re.search(r'\b(means|refers to|is defined as|definition)\b', lower)),
        "contains_formula": bool(re.search(r'[=\+\-\*/]\s*\d|ratio\s*=|formula|calculate', lower)),
        "contains_requirement": bool(re.search(r'\b(shall|must|is required to|are required to)\b', lower)),
        "contains_deadline": bool(re.search(r'\b(by\s+\w+\s+\d{4}|deadline|due date|no later than|within\s+\d+\s+(?:days|months|years))\b', lower)),
        "contains_assignment": bool(re.search(r'\b(responsible for|accountable|shall ensure|must ensure|is responsible)\b', lower)),
        "contains_parameter": bool(re.search(r'\b(threshold|limit|minimum|maximum|ratio|percentage|buffer|floor|ceiling)\b', lower)),
    }


def _infer_regulator_from_path(filepath: str) -> Dict[str, str]:
    """Infer regulator info from an ADLS file path.

    Args:
        filepath: ADLS blob path to scan for regulator acronyms.

    Returns:
        Dict with "regulator", "regulator_acronym", and "jurisdiction" keys.
    """
    upper_path = filepath.upper()
    for reg in KNOWN_REGULATORS:
        if reg["acronym"] in upper_path:
            return {
                "regulator": reg["name"],
                "regulator_acronym": reg["acronym"],
                "jurisdiction": reg["jurisdiction"],
            }
    return {"regulator": "", "regulator_acronym": "", "jurisdiction": ""}


def _infer_document_class_from_path(filepath: str) -> str:
    """Infer the document class from directory keywords in the path.

    Args:
        filepath: ADLS blob path to scan for document-class keywords.

    Returns:
        Document class string (e.g. "guideline", "policy") or "unknown".
    """
    lower_path = filepath.lower()
    mapping = {
        "guideline": "guideline",
        "advisory": "advisory",
        "rule": "rule",
        "policy": "policy",
        "policies": "policy",
        "procedure": "procedure",
        "procedures": "procedure",
        "standard": "standard",
        "standards": "standard",
        "circular": "circular",
        "notice": "notice",
        "consultation": "consultation",
        "framework": "framework",
    }
    for keyword, doc_class in mapping.items():
        if keyword in lower_path:
            return doc_class
    return "unknown"

# COMMAND ----------

# === NOVA Semantic Header & Prompt Rendering (Rules 1 & 3) ===


def build_semantic_header(meta: Dict) -> str:
    """Build a semantic header string prepended to chunk text before embedding (Rule 1).

    Produces a bracket-delimited header like
    [OSFI | B-20 | guideline | Capital Requirements > 3.2.1 | mandatory]
    so the vector captures document identity and structural signals.

    Args:
        meta: Dict of NOVA metadata fields for the chunk.

    Returns:
        Header string ending with newline, or "" if no fields are populated.
    """
    parts = []
    if meta.get("regulator_acronym"):
        parts.append(meta["regulator_acronym"])
    if meta.get("guideline_number"):
        parts.append(meta["guideline_number"])
    if meta.get("document_class"):
        parts.append(meta["document_class"])
    section_display = ""
    if meta.get("section_path"):
        section_display = meta["section_path"]
    if meta.get("section_number"):
        section_display = f"{section_display} > {meta['section_number']}" if section_display else meta["section_number"]
    if section_display:
        parts.append(section_display)
    if meta.get("normative_weight"):
        parts.append(meta["normative_weight"])
    if not parts:
        return ""
    return "[" + " | ".join(parts) + "]\n"


def render_chunk_for_prompt(chunk_text: str, meta: Dict) -> str:
    """Render a chunk with metadata headers for LLM prompt injection (Rule 3).

    Wraps the raw chunk text with relevant metadata so the LLM can
    reason about provenance and authority at inference time.

    Args:
        chunk_text: Raw text content of the chunk.
        meta: Dict of NOVA metadata fields for the chunk.

    Returns:
        Chunk text prefixed with a YAML-style metadata block, or
        the original text if no prompt-injected fields are populated.
    """
    source_type = meta.get("source_type", "auto")
    if source_type == "regulatory":
        prompt_fields = PROMPT_INJECTED_FIELDS_REGULATORY
    else:
        prompt_fields = PROMPT_INJECTED_FIELDS_INTERNAL

    header_lines = []
    for field in prompt_fields:
        val = meta.get(field)
        if val and val not in ("", [], None, False):
            if isinstance(val, list):
                val = "; ".join(str(v) for v in val)
            header_lines.append(f"{field}: {val}")

    if header_lines:
        header_block = "\n".join(header_lines)
        return f"---\n{header_block}\n---\n{chunk_text}"
    return chunk_text

# COMMAND ----------

# === NOVA Metadata Enrichment ===


def load_pre_extracted_metadata(filepath: str, container_client) -> Optional[Dict]:
    """Load pre-extracted metadata JSON from the silver/metadata/ layer.

    Args:
        filepath: Original blob path whose basename determines the JSON key.
        container_client: Azure ContainerClient for the storage container.

    Returns:
        Parsed metadata dict if found, else None.
    """
    basename = os.path.splitext(os.path.basename(filepath))[0]
    meta_blob_name = f"{SILVER_METADATA_PREFIX}{basename}.json"

    try:
        blob_client = container_client.get_blob_client(meta_blob_name)
        meta_bytes = blob_client.download_blob().readall()
        raw = json.loads(meta_bytes.decode("utf-8", errors="replace"))
        return raw.get("resolved_metadata", raw)
    except Exception:
        return None


def enrich_chunk_with_structural_metadata(doc: Document, doc_meta: Optional[Dict] = None) -> Document:
    """Merge pre-extracted document metadata and compute structural metadata.

    Operates on a single LangChain Document in-place, adding NOVA fields
    to doc.metadata.

    Args:
        doc: LangChain Document with at least page_content and basic metadata.
        doc_meta: Optional dict from load_pre_extracted_metadata().

    Returns:
        The same Document, enriched with NOVA metadata fields.
    """
    meta = doc.metadata
    text = doc.page_content or ""
    heading = meta.get("heading_path", "") or meta.get("Header 1", "")
    filepath = meta.get("file_path", "")

    # Merge pre-extracted doc-level metadata if available
    if doc_meta:
        for key in ("title", "short_title", "document_class", "regulator",
                     "regulator_acronym", "guideline_number", "jurisdiction",
                     "authority_class", "authority_level", "nova_tier",
                     "status", "effective_date_start", "effective_date_end",
                     "version_id", "version_label", "current_version_flag",
                     "sector", "doc_family_id", "business_owner",
                     "business_line", "audience", "approval_status",
                     "confidentiality", "source_type"):
            if doc_meta.get(key) and not meta.get(key):
                meta[key] = doc_meta[key]

    # Infer from path if not already set
    if not meta.get("regulator_acronym"):
        reg_info = _infer_regulator_from_path(filepath)
        meta.setdefault("regulator", reg_info["regulator"])
        meta.setdefault("regulator_acronym", reg_info["regulator_acronym"])
        meta.setdefault("jurisdiction", reg_info["jurisdiction"])

    if not meta.get("document_class"):
        meta["document_class"] = _infer_document_class_from_path(filepath)

    # Source type inference
    if not meta.get("source_type"):
        if BRONZE_EXTERNAL_PREFIX in filepath:
            meta["source_type"] = "regulatory"
        elif BRONZE_INTERNAL_PREFIX in filepath:
            meta["source_type"] = "internal"
        else:
            meta["source_type"] = "auto"

    # Structural metadata
    depth = meta.get("depth", 0)
    meta.setdefault("section_number", _extract_section_number(heading))
    meta.setdefault("structural_level", _infer_structural_level(depth, heading))
    meta["normative_weight"] = _classify_normative_weight(text)
    meta["paragraph_role"] = _classify_paragraph_role(text, heading)
    meta.setdefault("is_appendix", "appendix" in heading.lower() or "annex" in heading.lower())
    meta["cross_references"] = _extract_cross_references(text)

    # Content flags
    flags = _compute_content_flags(text)
    meta.update(flags)

    # BM25 text -- plain text without markup for lexical search
    meta["bm25_text"] = re.sub(r'[#\*\-\|>]', ' ', text).strip()

    # Citation anchor
    reg_acr = meta.get("regulator_acronym", "")
    gl_num = meta.get("guideline_number", "")
    sec_num = meta.get("section_number", "")
    if reg_acr and gl_num:
        meta["citation_anchor"] = f"{reg_acr}-{gl_num}" + (f"-{sec_num}" if sec_num else "")
    elif filepath:
        meta.setdefault("citation_anchor", os.path.basename(filepath))

    # Doc ID -- deterministic hash
    if not meta.get("doc_id"):
        meta["doc_id"] = hashlib.sha256(filepath.encode()).hexdigest()[:16]

    # Section path fallback
    meta.setdefault("section_path", heading)

    return doc

# COMMAND ----------

# === Regulatory Scraped JSON Parser ===


def process_regulatory_scraped_json(json_bytes: bytes, filepath: str) -> List[Document]:
    """Parse pre-scraped regulatory JSON into LangChain Document chunks.

    Accepts a dict with "title", "guideline_number", and "sections" keys,
    or a flat list of section objects. Each section becomes a Document
    with full NOVA metadata.

    Args:
        json_bytes: Raw JSON bytes from the blob.
        filepath: ADLS blob path (used for regulator/class inference).

    Returns:
        List of LangChain Document objects, one per section.
    """
    documents = []
    try:
        data = json.loads(json_bytes.decode("utf-8", errors="replace"))
    except json.JSONDecodeError as e:
        log_and_print(f"Invalid JSON in {filepath}: {e}", "error")
        return documents

    if isinstance(data, list):
        data = {"title": os.path.basename(filepath), "sections": data}

    title = data.get("title", os.path.basename(filepath))
    guideline_number = data.get("guideline_number", "")
    regulator_info = _infer_regulator_from_path(filepath)
    doc_class = _infer_document_class_from_path(filepath)

    sections = data.get("sections", [])
    if not sections and data.get("text"):
        sections = [{"heading": title, "text": data["text"]}]

    for idx, sec in enumerate(sections):
        heading = sec.get("heading", "")
        text = sec.get("text", "")
        if not text or not text.strip():
            continue

        meta = {
            "file_path": filepath,
            "source": os.path.basename(filepath),
            "filename": os.path.basename(filepath),
            "content_type": "text",
            "source_type": "regulatory",
            "title": title,
            "short_title": title[:80],
            "document_class": doc_class,
            "regulator": regulator_info.get("regulator", ""),
            "regulator_acronym": regulator_info.get("regulator_acronym", ""),
            "guideline_number": guideline_number,
            "jurisdiction": regulator_info.get("jurisdiction", ""),
            "section_path": heading,
            "section_number": _extract_section_number(heading),
            "structural_level": _infer_structural_level(sec.get("depth", 1), heading),
            "depth": sec.get("depth", 1),
            "normative_weight": _classify_normative_weight(text),
            "paragraph_role": _classify_paragraph_role(text, heading),
            "is_appendix": "appendix" in heading.lower() or "annex" in heading.lower(),
            "cross_references": _extract_cross_references(text),
            "chunk_index": idx,
            "total_chunks": len(sections),
        }
        meta.update(_compute_content_flags(text))

        doc = Document(page_content=text, metadata=meta)
        documents.append(doc)

    log_and_print(f"  Regulatory JSON parsed: {len(documents)} sections from {filepath}")
    return documents

# COMMAND ----------

# === Elasticsearch Connection & Index Management ===


def get_azure_credential():
    """Create Azure credential using client secret.

    Returns:
        ClientSecretCredential configured with tenant/client env vars.
    """
    return ClientSecretCredential(
        tenant_id=AZURE_TENANT_ID,
        client_id=AZURE_CLIENT_ID,
        client_secret=AZURE_CLIENT_SECRET,
    )


def get_blob_service_client():
    """Create BlobServiceClient with Azure credential.

    Returns:
        BlobServiceClient connected to the configured storage account.
    """
    credential = get_azure_credential()
    return BlobServiceClient(
        account_url=AZURE_STORAGE_ACCOUNT_URL,
        credential=credential,
    )


def get_es_client():
    """Create Elasticsearch client with auto-detected auth method.

    Supports cloud_id + api_key, cloud_id + basic_auth,
    host + basic_auth, and host + ca_fingerprint.

    Returns:
        Elasticsearch client instance (raises ConnectionError on failure).
    """
    es_kwargs = {
        "request_timeout": 60,
        "verify_certs": True,
        "retry_on_timeout": True,
        "max_retries": 3,
    }

    if ELASTICSEARCH_CLOUD_ID:
        es_kwargs["cloud_id"] = ELASTICSEARCH_CLOUD_ID
        if ELASTICSEARCH_API_KEY:
            es_kwargs["api_key"] = ELASTICSEARCH_API_KEY
        elif ELASTICSEARCH_BASIC_PASS:
            es_kwargs["basic_auth"] = (ELASTICSEARCH_BASIC_USER, ELASTICSEARCH_BASIC_PASS)
    elif ELASTICSEARCH_HOST:
        es_kwargs["hosts"] = [ELASTICSEARCH_HOST]
        if ELASTICSEARCH_BASIC_PASS:
            es_kwargs["basic_auth"] = (ELASTICSEARCH_BASIC_USER, ELASTICSEARCH_BASIC_PASS)
        if ELASTICSEARCH_CA_FINGERPRINT:
            es_kwargs["ssl_assert_fingerprint"] = ELASTICSEARCH_CA_FINGERPRINT

    es_client = Elasticsearch(**es_kwargs)

    if not es_client.ping():
        raise ConnectionError("Cannot reach Elasticsearch")

    log_and_print("Connected to Elasticsearch")
    return es_client


def create_es_vector_store(es_client, index_name=None):
    """Create Elasticsearch index with full NOVA vector mapping if it does not exist.

    Includes settings block, dense_vector (1024 dims, cosine), and all 40+ NOVA
    metadata fields mapped at top level for filter/faceted search.

    Args:
        es_client: Elasticsearch client instance.
        index_name: Target index name (defaults to INDEX_NAME).
    """
    index_name = index_name or INDEX_NAME

    if es_client.indices.exists(index=index_name):
        log_and_print(f"Index '{index_name}' already exists")
        return

    mapping = {
        "settings": {
            "number_of_shards": 1,
            "number_of_replicas": 1,
        },
        "mappings": {
            "properties": {
                # Core fields
                "doc_id": {"type": "keyword"},
                "chunk_id": {"type": "keyword"},
                "chunk_index": {"type": "integer"},
                "total_chunks": {"type": "integer"},
                "chunk_text": {"type": "text", "analyzer": "standard"},
                "source_file": {"type": "keyword"},
                "source_path": {
                    "type": "text",
                    "fields": {"keyword": {"type": "keyword"}},
                },
                "file_type": {"type": "keyword"},
                "title": {"type": "text", "fields": {"keyword": {"type": "keyword"}}},
                "dense_vector": {
                    "type": "dense_vector",
                    "dims": 1024,
                    "index": True,
                    "similarity": "cosine",
                },
                "ingestion_timestamp": {"type": "date"},
                "metadata.file_path_keyword": {"type": "keyword"},
                # Content metadata
                "document_type": {"type": "keyword"},
                "content_type": {"type": "keyword"},
                "has_tables": {"type": "boolean"},
                "has_images": {"type": "boolean"},
                "extraction_method": {"type": "keyword"},
                "word_count": {"type": "integer"},
                "char_count": {"type": "integer"},
                "page_number": {"type": "integer"},
                "heading_path": {"type": "text"},
                # ---- NOVA metadata fields (Rule 2 -- INDEX_FIELDS) ----
                "source_type": {"type": "keyword"},
                "short_title": {"type": "text", "fields": {"keyword": {"type": "keyword"}}},
                "document_class": {"type": "keyword"},
                "section_path": {"type": "text", "fields": {"keyword": {"type": "keyword"}}},
                "citation_anchor": {"type": "keyword"},
                "status": {"type": "keyword"},
                "effective_date_start": {"type": "date", "format": "yyyy-MM-dd||yyyy-MM||yyyy||epoch_millis", "ignore_malformed": True},
                "effective_date_end": {"type": "date", "format": "yyyy-MM-dd||yyyy-MM||yyyy||epoch_millis", "ignore_malformed": True},
                "jurisdiction": {"type": "keyword"},
                "authority_class": {"type": "keyword"},
                "authority_level": {"type": "keyword"},
                "nova_tier": {"type": "keyword"},
                "regulator": {"type": "text", "fields": {"keyword": {"type": "keyword"}}},
                "regulator_acronym": {"type": "keyword"},
                "guideline_number": {"type": "keyword"},
                "version_id": {"type": "keyword"},
                "version_label": {"type": "keyword"},
                "current_version_flag": {"type": "boolean"},
                "sector": {"type": "keyword"},
                "doc_family_id": {"type": "keyword"},
                "business_owner": {"type": "keyword"},
                "business_line": {"type": "keyword"},
                "audience": {"type": "keyword"},
                "approval_status": {"type": "keyword"},
                "confidentiality": {"type": "keyword"},
                "structural_level": {"type": "keyword"},
                "section_number": {"type": "keyword"},
                "depth": {"type": "integer"},
                "normative_weight": {"type": "keyword"},
                "paragraph_role": {"type": "keyword"},
                "is_appendix": {"type": "boolean"},
                "cross_references": {"type": "keyword"},
                "contains_definition": {"type": "boolean"},
                "contains_formula": {"type": "boolean"},
                "contains_requirement": {"type": "boolean"},
                "contains_deadline": {"type": "boolean"},
                "contains_assignment": {"type": "boolean"},
                "contains_parameter": {"type": "boolean"},
                "bm25_text": {"type": "text", "analyzer": "standard"},
            }
        }
    }

    es_client.indices.create(index=index_name, body=mapping)
    log_and_print(f"Created ES index '{index_name}' with NOVA vector mapping")

# COMMAND ----------

# === PGVector Dual-Store Support ===


def get_pg_conn():
    """Create a PostgreSQL connection for PGVector dual-store.

    Returns:
        psycopg2 connection, or None if PG_HOST is not set or connection fails.
    """
    if not PG_HOST:
        return None
    try:
        import psycopg2
        conn = psycopg2.connect(
            host=PG_HOST,
            port=PG_PORT,
            dbname=PG_DB,
            user=PG_USER,
            password=PG_PASSWORD,
        )
        conn.autocommit = False
        return conn
    except ImportError:
        log_and_print("psycopg2 not installed -- PGVector dual-store disabled", "warning")
        return None
    except Exception as e:
        log_and_print(f"PGVector connection failed: {e}", "warning")
        return None


def create_pgvector_table(conn):
    """Create the PGVector chunks table and indexes if they do not exist.

    Args:
        conn: psycopg2 connection with autocommit=False.
    """
    cur = conn.cursor()
    try:
        cur.execute("CREATE EXTENSION IF NOT EXISTS vector;")
        cur.execute(f"""
            CREATE TABLE IF NOT EXISTS {PG_TABLE} (
                chunk_id        TEXT PRIMARY KEY,
                doc_id          TEXT,
                chunk_index     INTEGER,
                total_chunks    INTEGER,
                chunk_text      TEXT,
                source_file     TEXT,
                source_path     TEXT,
                file_type       TEXT,
                title           TEXT,
                short_title     TEXT,
                source_type     TEXT,
                document_class  TEXT,
                section_path    TEXT,
                section_number  TEXT,
                citation_anchor TEXT,
                status          TEXT,
                effective_date_start TEXT,
                effective_date_end   TEXT,
                jurisdiction    TEXT,
                authority_class TEXT,
                authority_level TEXT,
                nova_tier       TEXT,
                regulator       TEXT,
                regulator_acronym TEXT,
                guideline_number TEXT,
                version_id      TEXT,
                version_label   TEXT,
                current_version_flag BOOLEAN DEFAULT TRUE,
                sector          TEXT,
                doc_family_id   TEXT,
                business_owner  TEXT,
                business_line   TEXT,
                audience        TEXT,
                approval_status TEXT,
                confidentiality TEXT,
                structural_level TEXT,
                depth           INTEGER DEFAULT 0,
                normative_weight TEXT,
                paragraph_role  TEXT,
                is_appendix     BOOLEAN DEFAULT FALSE,
                cross_references TEXT[],
                contains_definition BOOLEAN DEFAULT FALSE,
                contains_formula    BOOLEAN DEFAULT FALSE,
                contains_requirement BOOLEAN DEFAULT FALSE,
                contains_deadline   BOOLEAN DEFAULT FALSE,
                contains_assignment BOOLEAN DEFAULT FALSE,
                contains_parameter  BOOLEAN DEFAULT FALSE,
                bm25_text       TEXT,
                page_number     INTEGER,
                heading_path    TEXT,
                content_type    TEXT,
                ingestion_timestamp TIMESTAMPTZ DEFAULT NOW(),
                embedding       vector({EMBEDDING_DIMS})
            );
        """)
        cur.execute(f"CREATE INDEX IF NOT EXISTS idx_{PG_TABLE}_doc_id ON {PG_TABLE}(doc_id);")
        cur.execute(f"CREATE INDEX IF NOT EXISTS idx_{PG_TABLE}_regulator ON {PG_TABLE}(regulator_acronym);")
        conn.commit()
        log_and_print(f"PGVector table '{PG_TABLE}' ensured")
    except Exception as e:
        conn.rollback()
        log_and_print(f"Error creating PGVector table: {e}", "error")
    finally:
        cur.close()


def upsert_chunks_to_pgvector(conn, chunks: List[Document], embeddings: Optional[List] = None):
    """Upsert chunk documents into the PGVector table.

    Args:
        conn: psycopg2 connection.
        chunks: List of LangChain Document objects with NOVA metadata.
        embeddings: Optional list of embedding vectors (parallel to chunks).
    """
    if conn is None:
        return

    cur = conn.cursor()
    upserted = 0
    try:
        for i, doc in enumerate(chunks):
            meta = doc.metadata
            chunk_id = meta.get("chunk_id", "")
            if not chunk_id:
                raw = f"{meta.get('file_path', '')}::{i}"
                chunk_id = hashlib.sha256(raw.encode()).hexdigest()[:20]

            embedding_val = None
            if embeddings and i < len(embeddings):
                embedding_val = embeddings[i]

            cross_refs = meta.get("cross_references", [])
            if not isinstance(cross_refs, list):
                cross_refs = []

            cur.execute(f"""
                INSERT INTO {PG_TABLE} (
                    chunk_id, doc_id, chunk_index, total_chunks, chunk_text,
                    source_file, source_path, file_type, title, short_title,
                    source_type, document_class, section_path, section_number,
                    citation_anchor, status, effective_date_start, effective_date_end,
                    jurisdiction, authority_class, authority_level, nova_tier,
                    regulator, regulator_acronym, guideline_number,
                    version_id, version_label, current_version_flag,
                    sector, doc_family_id, business_owner, business_line,
                    audience, approval_status, confidentiality,
                    structural_level, depth, normative_weight, paragraph_role,
                    is_appendix, cross_references,
                    contains_definition, contains_formula, contains_requirement,
                    contains_deadline, contains_assignment, contains_parameter,
                    bm25_text, page_number, heading_path, content_type,
                    embedding
                ) VALUES (
                    %s, %s, %s, %s, %s, %s, %s, %s, %s, %s,
                    %s, %s, %s, %s, %s, %s, %s, %s, %s, %s,
                    %s, %s, %s, %s, %s, %s, %s, %s, %s, %s,
                    %s, %s, %s, %s, %s, %s, %s, %s, %s, %s,
                    %s, %s, %s, %s, %s, %s, %s, %s, %s, %s,
                    %s, %s
                )
                ON CONFLICT (chunk_id) DO UPDATE SET
                    chunk_text = EXCLUDED.chunk_text,
                    embedding = EXCLUDED.embedding,
                    normative_weight = EXCLUDED.normative_weight,
                    paragraph_role = EXCLUDED.paragraph_role,
                    cross_references = EXCLUDED.cross_references,
                    bm25_text = EXCLUDED.bm25_text,
                    ingestion_timestamp = NOW();
            """, (
                chunk_id, meta.get("doc_id", ""), meta.get("chunk_index", i),
                meta.get("total_chunks", 0), doc.page_content,
                meta.get("source", ""), meta.get("file_path", ""),
                meta.get("file_type", ""), meta.get("title", ""),
                meta.get("short_title", ""), meta.get("source_type", ""),
                meta.get("document_class", ""), meta.get("section_path", ""),
                meta.get("section_number", ""), meta.get("citation_anchor", ""),
                meta.get("status", ""), meta.get("effective_date_start", "") or None,
                meta.get("effective_date_end", "") or None, meta.get("jurisdiction", ""),
                meta.get("authority_class", ""), meta.get("authority_level", ""),
                meta.get("nova_tier", ""), meta.get("regulator", ""),
                meta.get("regulator_acronym", ""), meta.get("guideline_number", ""),
                meta.get("version_id", ""), meta.get("version_label", ""),
                meta.get("current_version_flag", True), meta.get("sector", ""),
                meta.get("doc_family_id", ""), meta.get("business_owner", ""),
                meta.get("business_line", ""), meta.get("audience", ""),
                meta.get("approval_status", ""), meta.get("confidentiality", ""),
                meta.get("structural_level", ""), meta.get("depth", 0),
                meta.get("normative_weight", ""), meta.get("paragraph_role", ""),
                meta.get("is_appendix", False), cross_refs,
                meta.get("contains_definition", False), meta.get("contains_formula", False),
                meta.get("contains_requirement", False), meta.get("contains_deadline", False),
                meta.get("contains_assignment", False), meta.get("contains_parameter", False),
                meta.get("bm25_text", ""), meta.get("page_number", 0),
                meta.get("heading_path", ""), meta.get("content_type", "text"),
                embedding_val,
            ))
            upserted += 1

        conn.commit()
        log_and_print(f"  PGVector: upserted {upserted} chunks")
    except Exception as e:
        conn.rollback()
        log_and_print(f"PGVector upsert error: {e}", "error")
    finally:
        cur.close()

# COMMAND ----------

# === ADLS File Operations ===


def fetch_adls_files_in_memory(token_manager=None):
    """Fetch files from ADLS into memory.

    Uses Azure credentials from environment variables to connect to ADLS.
    Separates JSON metadata files from document files.

    Args:
        token_manager: Optional OAuth token manager (unused, for interface compat).

    Returns:
        Tuple of (files, json_files) where:
            files: List of (filepath, filename, file_bytes) tuples for document files
            json_files: Dict of JSON metadata keyed by base filename
    """
    credential = get_azure_credential()
    blob_service_client = BlobServiceClient(
        account_url=AZURE_STORAGE_ACCOUNT_URL,
        credential=credential,
    )
    local_container_client = blob_service_client.get_container_client(container_name)

    files = []
    json_files = {}
    file_count = 0
    dir_count = 0

    for blob in local_container_client.list_blobs(name_starts_with=storage_folder):
        if "/" in blob.name:
            blob_client = local_container_client.get_blob_client(blob)
            file_bytes = blob_client.download_blob().readall()
            filepath = blob.name
            filename = os.path.basename(blob.name)
            file_count += 1
            print(f"\n  File #{file_count}: {filename}")

            if filename.endswith('.json'):
                base_filename = os.path.splitext(filename)[0]
                json_files[base_filename] = file_bytes
                log_and_print(f"  Found JSON metadata: {base_filename}")
            else:
                files.append((filepath, filename, file_bytes))
                log_and_print(f"  Fetched document: {filename}")
        else:
            dir_count += 1
            print(f"\n  Directory: {blob.name}")

    print(f"\n{'='*60}")
    print(f"Summary: Found {file_count} files and skipped {dir_count} directories")
    print(f"{'='*60}\n")

    return files, json_files

# COMMAND ----------

# === Check File Existence ===


def check_file_exists_in_index(es_client, index_name, file_path):
    """Check if a file_path already exists in the Elasticsearch index.

    Args:
        es_client: Elasticsearch client instance.
        index_name: Target index name to search.
        file_path: ADLS blob path to look up.

    Returns:
        True if the file is already indexed, False otherwise.
    """
    try:
        if not es_client.indices.exists(index=index_name):
            return False

        query = {
            "query": {
                "term": {
                    "source_path.keyword": file_path
                }
            },
            "size": 1
        }

        result = es_client.search(index=index_name, body=query)
        exists = result["hits"]["total"]["value"] > 0

        if exists:
            log_and_print(f"  File already exists in index: {file_path}")

        return exists

    except Exception as e:
        log_and_print(f"  Error checking file existence: {e}", "warning")
        return False

# COMMAND ----------

# === HTML Processing Helpers ===


def clean_html(html_content):
    """Parse and clean HTML content.

    Removes scripts, styles, navigation, footer, and other non-content elements.

    Args:
        html_content: Raw HTML string.

    Returns:
        BeautifulSoup object with cleaned content.
    """
    soup = BeautifulSoup(html_content, 'html.parser')

    unwanted_tags = ['script', 'style', 'nav', 'footer', 'header', 'aside',
                     'dialog', 'template', 'noscript', 'svg']
    for tag in soup.find_all(unwanted_tags):
        tag.decompose()

    unwanted_classes = [
        'navbar', 'nav', 'footer', 'sidebar', 'menu',
        'advertisement', 'ad', 'popup', 'modal', 'banner',
    ]
    for cls in unwanted_classes:
        for element in soup.find_all(class_=re.compile(cls, re.I)):
            element.decompose()

    return soup


def extract_text_from_html(html_content):
    """Extract readable text from HTML content.

    Removes non-content elements and preserves heading hierarchy with
    markdown-style formatting for headings.

    Args:
        html_content: Raw HTML string.

    Returns:
        Clean text suitable for chunking and embedding.
    """
    soup = clean_html(html_content)

    text_parts = []
    for element in soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'p', 'li', 'td', 'th']):
        if element.name and element.name.startswith('h'):
            level = int(element.name[1])
            text_parts.append(f"\n{'#' * level} {element.get_text(strip=True)}\n")
        else:
            text = element.get_text(strip=True)
            if text:
                text_parts.append(text)

    return '\n'.join(text_parts)


def process_html_images(html_content, html_dir=None, token_manager=None,
                        skip_duplicates=True):
    """Extract and process images from an HTML document using GPT Vision.

    Processes both local and remote images. For each image:
    1. Extract image bytes (from base64, local file, or URL)
    2. Optionally check for duplicates
    3. Process with GPT-4o Vision for content extraction

    Args:
        html_content: Raw HTML string.
        html_dir: Directory containing the HTML file (for resolving relative paths).
        token_manager: OAuth token manager for API calls.
        skip_duplicates: Whether to skip duplicate images based on hash.

    Returns:
        List of dicts with 'content' (vision analysis) and 'alt_text' keys.
    """
    if not html_content:
        return []

    soup = BeautifulSoup(html_content, 'html.parser')
    img_tags = soup.find_all('img')

    if not img_tags:
        return []

    image_results = []
    seen_hashes = set()

    for img_tag in img_tags:
        img_src = img_tag.get('src', '')
        alt_text = img_tag.get('alt', '')

        if not img_src:
            continue

        image_bytes = None

        if img_src.startswith('data:image'):
            try:
                _header, encoded = img_src.split(',', 1)
                image_bytes = base64.b64decode(encoded)
            except Exception as e:
                log_and_print(f"Error decoding base64 image: {e}", "warning")
                continue

        elif img_src.startswith(('http://', 'https://')) and httpx is not None:
            try:
                response = httpx.get(img_src, timeout=10.0)
                response.raise_for_status()
                image_bytes = response.content
            except Exception as e:
                log_and_print(f"Error downloading image from {img_src}: {e}", "warning")
                continue

        else:
            img_path = img_src
            if html_dir:
                img_path = os.path.join(html_dir, img_src)
            if os.path.exists(img_path):
                with open(img_path, 'rb') as f:
                    image_bytes = f.read()
            else:
                log_and_print(f"Image file not found: {img_path}", "warning")
                continue

        if not image_bytes:
            continue

        if skip_duplicates:
            image_hash = hashlib.md5(image_bytes).hexdigest()
            if image_hash in seen_hashes:
                continue
            seen_hashes.add(image_hash)

        if OCR_AVAILABLE and token_manager:
            try:
                ocr_processor = DocumentOCRProcessor(token_fetcher=token_manager)
                prompt = (
                    f"Analyze this image from an HTML document. "
                    f"Content: {alt_text if alt_text else 'No alt text provided'}. "
                    f"If chart/graph: describe type, extract data points and trends. "
                    f"If table: extract in markdown format. "
                    f"If text: extract visible text. Be precise."
                )
                ocr_result = ocr_processor.analyze_image(image_bytes, prompt, detail="high")
                image_results.append({
                    'content': ocr_result,
                    'alt_text': alt_text,
                })
            except Exception as e:
                log_and_print(f"Error processing image with OCR: {e}", "warning")

    return image_results

# COMMAND ----------

# === PDF Processing with OCR ===


def _process_pdf_with_ocr(file_content, filename, filepath, token_manager,
                           extended_metadata=None):
    """Process PDF using OCR to extract text, images, and tables.

    Falls back to PyMuPDF for basic text extraction if OCR fails.

    Args:
        file_content: Raw PDF content as bytes.
        filename: Original filename for metadata.
        filepath: Original filepath for metadata.
        token_manager: OAuth token manager for API calls.
        extended_metadata: Optional dictionary of extended metadata from JSON file.

    Returns:
        List of Document objects with extracted content.
    """
    if extended_metadata is None:
        extended_metadata = {}

    documents = []
    temp_file_path = None

    try:
        with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as temp_file:
            temp_file.write(file_content)
            temp_file_path = temp_file.name

        try:
            ocr = DocumentOCRProcessor(token_fetcher=token_manager)

            log_and_print(f"Extracting content with OCR from: {filename}")
            result = ocr.extract_from_pdf(
                pdf_path=temp_file_path,
                extract_images=True,
                extract_text=True,
                skip_duplicate_images=True,
                detail="high",
            )

            if result.get('text'):
                base_metadata = {
                    'file_path': filepath,
                    'source': filename,
                    'file_directory': os.path.dirname(filepath),
                    'filename': filename,
                    'content_type': 'text',
                }
                doc = Document(
                    page_content=result['text'],
                    metadata={**base_metadata, **extended_metadata},
                )
                documents.append(doc)

            for table in result.get('tables', []):
                if table.get('content') and table['content'].strip():
                    page_text = table.get('page_text', '')
                    if page_text and page_text.strip():
                        full_content = f"TABLE data with page content for additional context:\n{table['content']}\n\ntable content:\n{page_text}"
                    else:
                        full_content = f"TABLE\n{table['content']}"
                    doc = Document(
                        page_content=full_content,
                        metadata={
                            'file_path': filepath,
                            'source': filename,
                            'file_directory': os.path.dirname(filepath),
                            'filename': filename,
                            'content_type': 'table',
                            **extended_metadata,
                        },
                    )
                    documents.append(doc)

            CHART_CHAR_MIN = 50
            for image in result.get('images', []):
                image_content = image.get('content', '')
                if image_content and len(image_content.strip()) > CHART_CHAR_MIN:
                    page_text = image.get('page_text', '')
                    if page_text and page_text.strip():
                        full_content = f"IMAGE/chart content with page context for additional context:\n{image_content}\n\npage content:\n{page_text}"
                    else:
                        full_content = f"IMAGE/chart content:\n{image_content}"
                    doc = Document(
                        page_content=full_content,
                        metadata={
                            'file_path': filepath,
                            'source': filename,
                            'file_directory': os.path.dirname(filepath),
                            'filename': filename,
                            'content_type': 'image',
                            **extended_metadata,
                        },
                    )
                    documents.append(doc)

            log_and_print(f"PDF OCR extraction complete for {filename}: {len(documents)} document chunks")

        except Exception as ocr_error:
            log_and_print(f"OCR extraction failed for {filename}: {ocr_error}", "error")
            log_and_print(f"Falling back to PyMuPDF extraction")

            try:
                pdf_doc = fitz.open(temp_file_path)
                for page_num in range(len(pdf_doc)):
                    page = pdf_doc[page_num]
                    page_text = page.get_text("text")
                    if page_text and page_text.strip():
                        doc = Document(
                            page_content=page_text.strip(),
                            metadata={
                                'file_path': filepath,
                                'source': filename,
                                'file_directory': os.path.dirname(filepath),
                                'filename': filename,
                                'content_type': 'text',
                                'page': page_num + 1,
                                'total_pages': len(pdf_doc),
                                'extraction_method': 'pymupdf_fallback',
                                **extended_metadata,
                            },
                        )
                        documents.append(doc)
                pdf_doc.close()
            except Exception as fallback_error:
                log_and_print(f"Fallback PDF loading also failed: {fallback_error}", "error")
                return []

        return documents

    except Exception as e:
        log_and_print(f"Error processing PDF {filename}: {e}", "error")
        return []
    finally:
        if temp_file_path and os.path.exists(temp_file_path):
            os.unlink(temp_file_path)

# COMMAND ----------

# === DOCX Processing with OCR ===


def _process_docx_with_ocr(file_content, filename, filepath, token_manager,
                            extended_metadata=None):
    """Process DOCX using OCR to extract text, images, and tables.

    Falls back to python-docx for basic text extraction if OCR fails.

    Args:
        file_content: Raw DOCX content as bytes.
        filename: Original filename for metadata.
        filepath: Original filepath for metadata.
        token_manager: OAuth token manager for API calls.
        extended_metadata: Optional dictionary of extended metadata from JSON file.

    Returns:
        List of Document objects with extracted content.
    """
    if extended_metadata is None:
        extended_metadata = {}

    documents = []
    temp_file_path = None

    try:
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as temp_file:
            temp_file.write(file_content)
            temp_file_path = temp_file.name

        try:
            ocr = DocumentOCRProcessor(token_fetcher=token_manager)

            log_and_print(f"Extracting content with OCR from: {filename}")
            result = ocr.extract_from_docx(
                docx_path=temp_file_path,
                extract_images=True,
                extract_text=True,
                skip_duplicate_images=True,
                detail="high",
            )

            if result.get('text'):
                base_metadata = {
                    'file_path': filepath,
                    'source': filename,
                    'file_directory': os.path.dirname(filepath),
                    'filename': filename,
                    'content_type': 'text',
                }
                doc = Document(
                    page_content=result['text'],
                    metadata={**base_metadata, **extended_metadata},
                )
                documents.append(doc)

            for table in result.get('tables', []):
                if table.get('content') and table['content'].strip():
                    page_text = table.get('page_text', '')
                    if page_text and page_text.strip():
                        full_content = f"TABLE data with page content for additional context:\n{table['content']}\n\ntable content:\n{page_text}"
                    else:
                        full_content = f"TABLE\n{table['content']}"
                    doc = Document(
                        page_content=full_content,
                        metadata={
                            'file_path': filepath,
                            'source': filename,
                            'file_directory': os.path.dirname(filepath),
                            'filename': filename,
                            'content_type': 'table',
                            **extended_metadata,
                        },
                    )
                    documents.append(doc)

            CHART_CHAR_MIN = 50
            for image in result.get('images', []):
                image_content = image.get('content', '')
                if image_content and len(image_content.strip()) > CHART_CHAR_MIN:
                    page_text = image.get('page_text', '')
                    if page_text and page_text.strip():
                        full_content = f"IMAGE/chart content with page context for additional context:\n{image_content}\n\npage content:\n{page_text}"
                    else:
                        full_content = f"IMAGE/chart content:\n{image_content}"
                    doc = Document(
                        page_content=full_content,
                        metadata={
                            'file_path': filepath,
                            'source': filename,
                            'file_directory': os.path.dirname(filepath),
                            'filename': filename,
                            'content_type': 'image',
                            **extended_metadata,
                        },
                    )
                    documents.append(doc)

            log_and_print(f"DOCX OCR extraction complete for {filename}: {len(documents)} document chunks")

        except Exception as ocr_error:
            log_and_print(f"OCR extraction failed for DOCX {filename}: {ocr_error}", "error")

            try:
                import docx as python_docx
                docx_file = io.BytesIO(file_content)
                doc_obj = python_docx.Document(docx_file)

                full_text = []
                for para in doc_obj.paragraphs:
                    if para.text.strip():
                        full_text.append(para.text.strip())

                if full_text:
                    doc = Document(
                        page_content='\n'.join(full_text),
                        metadata={
                            'file_path': filepath,
                            'source': filename,
                            'file_directory': os.path.dirname(filepath),
                            'filename': filename,
                            'content_type': 'text',
                            'extraction_method': 'python_docx_fallback',
                            **extended_metadata,
                        },
                    )
                    documents.append(doc)

                for table_idx, table in enumerate(doc_obj.tables):
                    table_text = ""
                    for row in table.rows:
                        row_cells = [cell.text.strip() for cell in row.cells]
                        table_text += " | ".join(row_cells) + "\n"
                    if table_text.strip():
                        doc = Document(
                            page_content=f"TABLE\n{table_text}",
                            metadata={
                                'file_path': filepath,
                                'source': filename,
                                'filename': filename,
                                'content_type': 'table',
                                'table_index': table_idx,
                                'extraction_method': 'python_docx_fallback',
                            },
                        )
                        documents.append(doc)

            except Exception as fallback_error:
                log_and_print(f"Fallback DOCX loading also failed: {fallback_error}", "error")
                return []

        return documents

    except Exception as e:
        log_and_print(f"Error processing DOCX {filename}: {e}", "error")
        return []
    finally:
        if temp_file_path and os.path.exists(temp_file_path):
            os.unlink(temp_file_path)

# COMMAND ----------

# === Text Chunking ===


def create_text_chunks_with_headings(text, chunk_size=None, chunk_overlap=None):
    """Divide text content into chunks using RecursiveCharacterTextSplitter.

    Tables and images should be passed separately and not chunked here.

    Args:
        text: Raw text to split into chunks.
        chunk_size: Maximum characters per chunk (default from CHUNK_SIZE env var).
        chunk_overlap: Number of characters to overlap between chunks.

    Returns:
        List of text chunk strings.
    """
    if chunk_size is None:
        chunk_size = CHUNK_SIZE
    if chunk_overlap is None:
        chunk_overlap = CHUNK_OVERLAP

    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        add_start_index=True,
    )

    return text_splitter.split_text(text)

# COMMAND ----------

# === HTML Processing with OCR ===


def process_html_with_ocr(file_content, filename, filepath, token_manager,
                           extended_metadata=None):
    """Process HTML with text extraction and GPT-4o Vision image analysis.

    Args:
        file_content: Raw HTML content as bytes.
        filename: Original filename for metadata.
        filepath: Original filepath for metadata.
        token_manager: OAuth token manager for API calls.
        extended_metadata: Optional dictionary of extended metadata from JSON file.

    Returns:
        List of Document objects.
    """
    if extended_metadata is None:
        extended_metadata = {}

    html_content = file_content.decode('utf-8', errors='replace')

    text_content = extract_text_from_html(html_content)

    html_dir = os.path.dirname(filepath)
    image_results = process_html_images(
        html_content,
        html_dir=html_dir,
        token_manager=token_manager,
        skip_duplicates=True,
    )

    documents = []

    if text_content and text_content.strip():
        doc = Document(
            page_content=text_content,
            metadata={
                'file_path': filepath,
                'source': filename,
                'file_directory': os.path.dirname(filepath),
                'filename': filename,
                'content_type': 'text',
                **extended_metadata,
            },
        )
        documents.append(doc)

    for img_result in image_results:
        if img_result.get('content') and len(img_result['content'].strip()) > 50:
            doc = Document(
                page_content=f"IMAGE/chart content:\n{img_result['content']}",
                metadata={
                    'file_path': filepath,
                    'source': filename,
                    'filename': filename,
                    'content_type': 'image',
                    'alt_text': img_result.get('alt_text', ''),
                },
            )
            documents.append(doc)

    log_and_print(f"HTML processing complete: {len(documents)} document chunks")
    return documents

# COMMAND ----------

# === Main Pipeline - load_and_split_documents() ===


def load_and_split_documents(files_to_process, json_files, token_manager,
                              path_category="auto"):
    """Load and split documents from ADLS into chunks for Elasticsearch indexing.

    For each file, detects file type and processes accordingly:
    - PDF/DOCX with OCR when available
    - HTML with text extraction and image analysis
    - TXT/MD loaded directly
    - JSON parsed as regulatory scraped content

    After extraction, text documents are chunked and all chunks are enriched
    with NOVA structural metadata.

    Args:
        files_to_process: List of tuples (filepath, filename, file_bytes).
        json_files: Dict of JSON metadata keyed by base filename.
        token_manager: OAuth token manager for API calls.
        path_category: Category of files being processed (regulatory/internal/auto).

    Returns:
        List of Document chunks ready for Elasticsearch ingestion.
    """
    all_documents = []
    failed_files = []

    for idx, (filepath, filename, file_content) in enumerate(files_to_process):
        log_and_print(f"\nProcessing file {idx + 1}/{len(files_to_process)}: {filename}")
        log_and_print(f"  Path category: {path_category}")

        try:
            base_name = os.path.splitext(filename)[0]
            extended_metadata = {}

            if base_name in json_files:
                try:
                    json_content = json_files[base_name]
                    if isinstance(json_content, bytes):
                        json_content = json_content.decode('utf-8')
                    extended_metadata = json.loads(json_content)
                    log_and_print(f"  Found JSON metadata for {filename}")
                except Exception as e:
                    log_and_print(f"  Error parsing JSON metadata: {e}", "warning")

            ext = os.path.splitext(filename)[1].lower()
            docs = []

            if ext == '.pdf':
                docs = _process_pdf_with_ocr(
                    file_content, filename, filepath, token_manager,
                    extended_metadata=extended_metadata,
                )
            elif ext == '.docx':
                docs = _process_docx_with_ocr(
                    file_content, filename, filepath, token_manager,
                    extended_metadata=extended_metadata,
                )
            elif ext in ('.html', '.htm'):
                docs = process_html_with_ocr(
                    file_content, filename, filepath, token_manager,
                    extended_metadata=extended_metadata,
                )
            elif ext == '.json':
                docs = process_regulatory_scraped_json(file_content, filepath)
            elif ext in ('.txt', '.md'):
                content = file_content.decode('utf-8', errors='replace')
                doc = Document(
                    page_content=content,
                    metadata={
                        'file_path': filepath,
                        'source': filename,
                        'filename': filename,
                        'content_type': 'text',
                        **extended_metadata,
                    },
                )
                docs = [doc]
            else:
                log_and_print(f"  Skipping unsupported file type: {ext}", "warning")
                continue

            if docs:
                all_documents.extend(docs)
                log_and_print(f"  Extracted {len(docs)} document chunks from {filename}")
            else:
                log_and_print(f"  No content extracted from {filename}", "warning")

        except Exception as e:
            log_and_print(f"  Error processing {filename}: {e}", "error")
            failed_files.append((filename, str(e)))

    if failed_files:
        log_and_print(f"\nFailed files ({len(failed_files)}):")
        for fname, err in failed_files:
            log_and_print(f"  - {fname}: {err}", "error")

    # Chunk text documents; keep tables and images as single chunks
    text_docs = [d for d in all_documents if d.metadata.get('content_type') == 'text']
    non_text_docs = [d for d in all_documents if d.metadata.get('content_type') != 'text']

    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=CHUNK_SIZE,
        chunk_overlap=CHUNK_OVERLAP,
        add_start_index=True,
    )

    # Use markdown header splitting for markdown-like content
    headers_to_split_on = [
        ("#", "Header 1"),
        ("##", "Header 2"),
        ("###", "Header 3"),
    ]
    md_splitter = MarkdownHeaderTextSplitter(headers_to_split_on=headers_to_split_on)

    text_chunks = []
    for doc in text_docs:
        content = doc.page_content
        if not content or not content.strip():
            continue

        if content.strip().startswith("#") or "\n#" in content:
            try:
                md_chunks = md_splitter.split_text(content)
                for chunk in md_chunks:
                    chunk.metadata.update(doc.metadata)
                text_chunks.extend(md_chunks)
                continue
            except Exception:
                pass

        chunks = text_splitter.split_documents([doc])
        text_chunks.extend(chunks)

    log_and_print(f"Chunked {len(text_docs)} text documents into {len(text_chunks)} chunks")

    all_chunks = text_chunks + non_text_docs
    log_and_print(f"Total chunks (text + tables + images): {len(all_chunks)}")

    # NOVA: Enrich all chunks with structural metadata
    for doc in all_chunks:
        base = os.path.splitext(os.path.basename(doc.metadata.get("source", "")))[0]
        file_meta = json_files.get(base)
        if isinstance(file_meta, (bytes, str)):
            try:
                file_meta = json.loads(file_meta if isinstance(file_meta, str) else file_meta.decode("utf-8"))
            except Exception:
                file_meta = None
        enrich_chunk_with_structural_metadata(doc, doc_meta=file_meta)

    return all_chunks

# COMMAND ----------

# === Main Entry Point ===


def main():
    """Main entry point for the NOVA ingestion pipeline.

    Steps:
        1. Initialize OAuth token manager
        2. Fetch files from ADLS
        3. Connect to Elasticsearch
        4. Create/verify index with NOVA mapping
        5. Check existing documents
        6. Process and chunk documents
        6b. Apply Rule 1 semantic headers and assign chunk_ids
        7. Generate embeddings
        8. Ingest into Elasticsearch (NOVA INDEX_FIELDS at top level)
        9. PGVector dual-store upsert
    """
    log_and_print("=" * 60)
    log_and_print("STARTING NOVA DOCUMENT INGESTION PIPELINE")
    log_and_print("=" * 60)

    # Step 1: Initialize token manager
    token_manager = None
    if OCR_AVAILABLE:
        try:
            token_manager = OAuthTokenManager()
            log_and_print("Step 1: Token manager initialized successfully")
        except Exception as e:
            log_and_print(f"Step 1: Token manager init failed: {e}", "warning")
    else:
        log_and_print("Step 1: OCR not available -- skipping token manager")

    # Step 2: Fetch files from ADLS
    log_and_print("Step 2: Fetching files from ADLS")
    files_to_process, json_files = fetch_adls_files_in_memory(token_manager)

    if not files_to_process:
        log_and_print("No files found to process. Exiting.", "warning")
        return

    log_and_print(f"  Found {len(files_to_process)} document files and {len(json_files)} JSON metadata files")

    # Step 3: Connect to Elasticsearch
    log_and_print("Step 3: Connecting to Elasticsearch")
    es_client = get_es_client()

    # Step 4: Create/verify index
    index_name = INDEX_NAME
    log_and_print(f"Step 4: Creating/verifying index: {index_name}")
    create_es_vector_store(es_client, index_name)

    # Step 5: Check which files already exist
    log_and_print("Step 5: Checking which documents already exist")
    files_to_skip = []
    checked_files = set()

    for filepath, filename, _ in files_to_process:
        if filepath not in checked_files:
            checked_files.add(filepath)
            if check_file_exists_in_index(es_client, index_name, filepath):
                files_to_skip.append(filepath)

    if files_to_skip:
        filtered = [(fp, fn, fc) for fp, fn, fc in files_to_process if fp not in files_to_skip]
        log_and_print(f"  Filtered: {len(files_to_process)} -> {len(filtered)} files (skipped {len(files_to_skip)})")
        files_to_process = filtered

    if not files_to_process:
        log_and_print("All files already indexed. Nothing to ingest.")
        return

    # Step 6: Process and chunk documents
    log_and_print("Step 6: Processing and chunking documents")

    pdf_files = []
    docx_files = []
    html_files = []
    txt_files = []

    for filepath, filename, file_content in files_to_process:
        ext = os.path.splitext(filename)[1].lower()
        if ext == '.pdf':
            pdf_files.append((filepath, filename, file_content))
        elif ext == '.docx':
            docx_files.append((filepath, filename, file_content))
        elif ext in ('.html', '.htm'):
            html_files.append((filepath, filename, file_content))
        elif ext in ('.txt', '.md', '.json'):
            txt_files.append((filepath, filename, file_content))

    all_chunked_docs = []

    if pdf_files:
        log_and_print(f"  Processing {len(pdf_files)} PDF files...")
        pdf_chunks = load_and_split_documents(
            pdf_files, json_files, token_manager, path_category="regulatory",
        )
        all_chunked_docs.extend(pdf_chunks)

    if docx_files:
        log_and_print(f"  Processing {len(docx_files)} DOCX files...")
        docx_chunks = load_and_split_documents(
            docx_files, json_files, token_manager, path_category="internal_raw",
        )
        all_chunked_docs.extend(docx_chunks)

    if html_files:
        log_and_print(f"  Processing {len(html_files)} HTML files...")
        html_chunks = load_and_split_documents(
            html_files, json_files, token_manager, path_category="auto",
        )
        all_chunked_docs.extend(html_chunks)

    if txt_files:
        log_and_print(f"  Processing {len(txt_files)} text/JSON files...")
        txt_chunks = load_and_split_documents(
            txt_files, json_files, token_manager, path_category="auto",
        )
        all_chunked_docs.extend(txt_chunks)

    if not all_chunked_docs:
        log_and_print("No chunks created. Exiting.", "warning")
        return

    # Step 6b: Apply NOVA Rule 1 - Semantic Headers + assign chunk_ids
    log_and_print("Step 6b: Applying NOVA Rule 1 - Semantic headers and chunk IDs")
    for i, doc in enumerate(all_chunked_docs):
        header = build_semantic_header(doc.metadata)
        if header:
            doc.page_content = header + doc.page_content
        raw_id = f"{doc.metadata.get('file_path', '')}::{i}"
        doc.metadata["chunk_id"] = hashlib.sha256(raw_id.encode()).hexdigest()[:20]
        doc.metadata["chunk_index"] = i
        doc.metadata["total_chunks"] = len(all_chunked_docs)

    # Step 7: Generate embeddings
    log_and_print("Step 7: Generating embeddings")
    from openai import OpenAI
    openai_client = OpenAI()  # Uses OPENAI_API_KEY env var

    all_embeddings = []
    EMBED_BATCH_SIZE = 100
    texts_to_embed = [doc.page_content for doc in all_chunked_docs]

    for i in range(0, len(texts_to_embed), EMBED_BATCH_SIZE):
        batch_texts = texts_to_embed[i:i + EMBED_BATCH_SIZE]
        try:
            response = openai_client.embeddings.create(
                model=EMBEDDING_MODEL,
                input=batch_texts,
                dimensions=EMBEDDING_DIMS,
            )
            batch_embeddings = [item.embedding for item in response.data]
            all_embeddings.extend(batch_embeddings)
            log_and_print(f"  Embedded batch {i // EMBED_BATCH_SIZE + 1}: {len(batch_embeddings)} chunks")
        except Exception as e:
            log_and_print(f"  Embedding batch failed: {e}", "error")
            # Fill with None for failed batches
            all_embeddings.extend([None] * len(batch_texts))

    log_and_print(f"  Total embeddings generated: {sum(1 for e in all_embeddings if e is not None)}")

    # Step 8: Ingest into Elasticsearch
    log_and_print("Step 8: Ingesting into Elasticsearch")

    BATCH_SIZE = 50
    total_count = 0
    total_batches = (len(all_chunked_docs) + BATCH_SIZE - 1) // BATCH_SIZE

    for batch_idx in range(0, len(all_chunked_docs), BATCH_SIZE):
        batch = all_chunked_docs[batch_idx:batch_idx + BATCH_SIZE]
        current_batch = batch_idx // BATCH_SIZE + 1

        log_and_print(f"  Processing batch {current_batch}/{total_batches} ({len(batch)} chunks)")

        actions = []
        for j, chunk in enumerate(batch):
            source_doc = {
                "chunk_text": chunk.page_content,
                "source_file": chunk.metadata.get("source", ""),
                "source_path": chunk.metadata.get("file_path", ""),
                "file_type": os.path.splitext(chunk.metadata.get("source", ""))[1],
                "ingestion_timestamp": datetime.datetime.utcnow().isoformat(),
            }
            # Add all NOVA index fields at top level
            for field in INDEX_FIELDS:
                val = chunk.metadata.get(field)
                if val is not None:
                    source_doc[field] = val
            # Add remaining metadata fields
            for key in ("doc_id", "chunk_id", "chunk_index", "total_chunks",
                        "title", "page_number", "heading_path", "content_type",
                        "has_tables", "has_images", "extraction_method",
                        "word_count", "char_count"):
                val = chunk.metadata.get(key)
                if val is not None:
                    source_doc[key] = val

            # Add embedding vector
            embed_idx = batch_idx + j
            if embed_idx < len(all_embeddings) and all_embeddings[embed_idx] is not None:
                source_doc["dense_vector"] = all_embeddings[embed_idx]

            action = {
                "_index": index_name,
                "_source": source_doc,
            }
            if chunk.metadata.get("chunk_id"):
                action["_id"] = chunk.metadata["chunk_id"]
            actions.append(action)

        try:
            helpers.bulk(es_client, actions)
            total_count += len(batch)
            log_and_print(f"  Batch {current_batch} ingested successfully ({total_count} total)")
        except Exception as e:
            log_and_print(f"  Error ingesting batch {current_batch}: {e}", "error")

    # Step 9: PGVector Dual-Store (optional)
    log_and_print("Step 9: PGVector dual-store upsert")
    pg_conn = get_pg_conn()
    if pg_conn:
        try:
            create_pgvector_table(pg_conn)
            upsert_chunks_to_pgvector(pg_conn, all_chunked_docs, embeddings=all_embeddings)
            log_and_print("  PGVector dual-store upsert complete")
        except Exception as e:
            log_and_print(f"  PGVector upsert failed (non-fatal): {e}", "warning")
        finally:
            try:
                pg_conn.close()
            except Exception:
                pass
    else:
        log_and_print("  PGVector not configured (PG_HOST not set) -- skipping dual-store")

    # Summary
    log_and_print(f"\n{'='*60}")
    log_and_print("INGESTION PIPELINE COMPLETE")
    log_and_print(f"{'='*60}")

    try:
        result = es_client.count(index=index_name)
        result_docs = result.get("count", 0)
        log_and_print(f"  Total documents in index '{index_name}': {result_docs}")
    except Exception:
        pass

    log_and_print(f"  Total chunks ingested this run: {total_count}")
    log_and_print(f"  Files processed: {len(checked_files)}")
    log_and_print(f"  Files skipped (already indexed): {len(files_to_skip)}")
    log_and_print("Pipeline complete!")

# COMMAND ----------

# === Azure Functions Entry Point (alternative) ===

if AZURE_FUNCTIONS_AVAILABLE:
    app = func.FunctionApp()

    @app.function_name(name="IngestADLSFiles")
    @app.route(route="ingest", methods=["POST", "GET"])
    def azure_func_main(req: func.HttpRequest) -> func.HttpResponse:
        """Azure Functions HTTP trigger for the ingestion pipeline.

        Args:
            req: Azure Functions HttpRequest (POST or GET).

        Returns:
            HttpResponse with JSON status body (200 on success, 500 on error).
        """
        log_and_print("Received request to ingest ADLS files")

        try:
            main()
            return func.HttpResponse(
                json.dumps({"status": "success", "message": "Ingestion complete"}),
                status_code=200,
                mimetype="application/json",
            )
        except Exception as e:
            log_and_print(f"Ingestion failed: {e}", "error")
            traceback.print_exc()
            return func.HttpResponse(
                json.dumps({"status": "error", "message": str(e)}),
                status_code=500,
                mimetype="application/json",
            )

# COMMAND ----------

if __name__ == "__main__":
    main()
