# Databricks notebook source
# COMMAND ----------
# MAGIC %md
# MAGIC # NOVA Shared Utilities
# MAGIC
# MAGIC Shared models, helpers, parsing, canonicalization, embeddings, and storage functions
# MAGIC used by the ingestion and retrieval notebooks.
# MAGIC
# MAGIC **Key design principle**: the metadata spec (`01_metadata_spec.py`) is the single source
# MAGIC of truth for which fields are embedded (Rule 1), indexed (Rule 2), prompt-injected
# MAGIC (Rule 3), or operational-only. This module imports the spec and uses it to drive
# MAGIC `semantic_header()`, `build_chunk_docs()`, and `render_hit_for_prompt()`.

# COMMAND ----------
# MAGIC %pip install -q azure-identity azure-storage-file-datalake azure-ai-documentintelligence \
# MAGIC   elasticsearch psycopg[binary] pgvector openai beautifulsoup4 python-docx

# COMMAND ----------
from __future__ import annotations

import io
import json
import hashlib
import os
import re
from dataclasses import dataclass, asdict, field
from datetime import datetime, timezone
from importlib.util import module_from_spec, spec_from_file_location
from typing import Any, Dict, Iterable, List, Optional

from bs4 import BeautifulSoup
from docx import Document as DocxDocument
from azure.identity import DefaultAzureCredential
from azure.storage.filedatalake import DataLakeServiceClient
from azure.ai.documentintelligence import DocumentIntelligenceClient
from elasticsearch import Elasticsearch
from elasticsearch.helpers import bulk
from openai import OpenAI
import psycopg
from pgvector.psycopg import register_vector

# COMMAND ----------
# MAGIC %md
# MAGIC ## Load metadata spec (single source of truth)

# COMMAND ----------
_SPEC_PATH = os.path.join(os.path.dirname(os.path.abspath(__file__)), "01_metadata_spec.py")
_loader_spec = spec_from_file_location("metadata_spec", _SPEC_PATH)
if _loader_spec is None or _loader_spec.loader is None:
    raise RuntimeError(
        f"Unable to load 01_metadata_spec.py from {_SPEC_PATH}. "
        "Ensure it is in the same directory as this notebook."
    )
metadata_spec = module_from_spec(_loader_spec)
_loader_spec.loader.exec_module(metadata_spec)

# Pre-compute specs for both document types at import time.
REGULATORY_SPEC: Dict[str, List[str]] = metadata_spec.build_spec("regulatory")
INTERNAL_SPEC: Dict[str, List[str]] = metadata_spec.build_spec("internal")
UNIT_LEVEL_INDEX_FLAGS: List[str] = metadata_spec.UNIT_LEVEL_INDEX_FLAGS

# COMMAND ----------
# MAGIC %md
# MAGIC ## Runtime configuration
# MAGIC
# MAGIC These widgets make the notebook portable across dev / test / prod workspaces.

# COMMAND ----------
try:
    dbutils.widgets.text("adls_account_url", "")
    dbutils.widgets.text("adls_file_system", "nova-docs")
    dbutils.widgets.text("doc_intel_endpoint", "")
    dbutils.widgets.text("elastic_url", "")
    dbutils.widgets.text("chunk_index_name", "nova_chunks_v1")
    dbutils.widgets.text("embed_model", "text-embedding-3-large")
    dbutils.widgets.text("embed_dimensions", "1024")
    dbutils.widgets.text("pgvector_host", "")
    dbutils.widgets.text("pgvector_port", "5432")
    dbutils.widgets.text("pgvector_db", "nova")
    dbutils.widgets.text("pgvector_user", "")
    dbutils.widgets.text("adls_secret_scope", "")
    dbutils.widgets.text("openai_secret_scope", "")
    dbutils.widgets.text("elastic_secret_scope", "")
    dbutils.widgets.text("pgvector_secret_scope", "")
except Exception:
    pass

# COMMAND ----------
def widget(name: str, default: str = "") -> str:
    try:
        value = dbutils.widgets.get(name)
        return value if value is not None else default
    except Exception:
        return default


def get_secret(scope_widget: str, key: str, fallback_env: str = "") -> str:
    scope = widget(scope_widget, "")
    if scope:
        return dbutils.secrets.get(scope=scope, key=key)
    return os.environ.get(fallback_env or key.upper(), "")


ADLS_ACCOUNT_URL = widget("adls_account_url")
ADLS_FILE_SYSTEM = widget("adls_file_system", "nova-docs")
DOC_INTEL_ENDPOINT = widget("doc_intel_endpoint")
ELASTIC_URL = widget("elastic_url")
CHUNK_INDEX_NAME = widget("chunk_index_name", "nova_chunks_v1")
EMBED_MODEL = widget("embed_model", "text-embedding-3-large")
EMBED_DIMENSIONS = int(widget("embed_dimensions", "1024"))
PGVECTOR_HOST = widget("pgvector_host")
PGVECTOR_PORT = int(widget("pgvector_port", "5432"))
PGVECTOR_DB = widget("pgvector_db", "nova")
PGVECTOR_USER = widget("pgvector_user")

OPENAI_API_KEY = get_secret("openai_secret_scope", "openai_api_key", "OPENAI_API_KEY")
ELASTIC_API_KEY = get_secret("elastic_secret_scope", "elastic_api_key", "ELASTIC_API_KEY")
PGVECTOR_PASSWORD = get_secret("pgvector_secret_scope", "pgvector_password", "PGVECTOR_PASSWORD")

# COMMAND ----------
# MAGIC %md
# MAGIC ## Models

# COMMAND ----------
@dataclass
class DocumentRegistryRow:
    doc_id: str
    title: Optional[str]
    source_type: str
    raw_path: str
    raw_file_type: str
    source_system: Optional[str]
    source_url: Optional[str]
    sha256: str
    ingested_at: str
    parser_version: Optional[str] = None
    extraction_status: str = "pending"


@dataclass
class CanonicalUnit:
    unit_id: str
    unit_type: str
    heading_path: list[str]
    section_path: str
    text: Optional[str]
    citation_anchor: Optional[str] = None
    page_start: Optional[int] = None
    page_end: Optional[int] = None
    table_id: Optional[str] = None
    paragraph_anchor: Optional[str] = None
    contains_definition: bool = False
    contains_formula: bool = False
    contains_deadline: bool = False
    contains_requirement: bool = False
    contains_parameter: bool = False
    contains_assignment: bool = False


@dataclass
class CanonicalDocument:
    doc_id: str
    doc_type: str  # "regulatory" or "internal"
    title: str
    source_type: str
    document_class: str
    raw_path: str
    canonical_path: str
    source_system: Optional[str] = None
    source_url: Optional[str] = None
    short_title: Optional[str] = None
    version: Optional[str] = None
    status: Optional[str] = None
    effective_date_start: Optional[str] = None
    effective_date_end: Optional[str] = None
    jurisdiction: Optional[str] = None
    authority_class: Optional[str] = None
    authority_level: Optional[int] = None
    nova_tier: Optional[int] = None
    confidentiality: Optional[str] = None
    approval_status: Optional[str] = None
    business_owner: Optional[str] = None
    business_line: Optional[str] = None
    legal_entity: Optional[str] = None
    audience: Optional[str] = None
    regulator: Optional[str] = None
    regulator_acronym: Optional[str] = None
    doc_family_id: Optional[str] = None
    version_id: Optional[str] = None
    version_label: Optional[str] = None
    version_sort_key: Optional[str] = None
    guideline_number: Optional[str] = None
    current_version_flag: Optional[str] = None
    sector: Optional[str] = None
    supersedes_doc_id: Optional[str] = None
    superseded_by_doc_id: Optional[str] = None
    document_owner: Optional[str] = None
    approval_date: Optional[str] = None
    review_date: Optional[str] = None
    next_review_date: Optional[str] = None
    function: Optional[str] = None
    related_doc_ids: list[str] = field(default_factory=list)
    sha256: Optional[str] = None
    parser_version: Optional[str] = None
    quality_score: Optional[float] = None
    quality_flags: list[str] = field(default_factory=list)
    units: list[CanonicalUnit] = field(default_factory=list)

# COMMAND ----------
# MAGIC %md
# MAGIC ## Client helpers

# COMMAND ----------
def utc_now_iso() -> str:
    return datetime.now(timezone.utc).isoformat()


def sha256_bytes(data: bytes) -> str:
    return hashlib.sha256(data).hexdigest()


def get_credential() -> DefaultAzureCredential:
    return DefaultAzureCredential()


def get_adls_client() -> DataLakeServiceClient:
    return DataLakeServiceClient(account_url=ADLS_ACCOUNT_URL, credential=get_credential())


def get_fs_client():
    return get_adls_client().get_file_system_client(ADLS_FILE_SYSTEM)


def read_adls_bytes(path: str) -> bytes:
    return get_fs_client().get_file_client(path).download_file().readall()


def write_adls_json(path: str, payload: dict[str, Any]) -> None:
    data = json.dumps(payload, ensure_ascii=False, indent=2).encode("utf-8")
    get_fs_client().get_file_client(path).upload_data(data, overwrite=True)


def get_doc_intelligence_client() -> DocumentIntelligenceClient:
    return DocumentIntelligenceClient(DOC_INTEL_ENDPOINT, get_credential())


def get_elastic_client() -> Elasticsearch:
    return Elasticsearch(ELASTIC_URL, api_key=ELASTIC_API_KEY)


def get_openai_client() -> OpenAI:
    return OpenAI(api_key=OPENAI_API_KEY)


def get_pg_conn(autocommit: bool = True):
    conn = psycopg.connect(
        host=PGVECTOR_HOST,
        port=PGVECTOR_PORT,
        dbname=PGVECTOR_DB,
        user=PGVECTOR_USER,
        password=PGVECTOR_PASSWORD,
        autocommit=autocommit,
    )
    register_vector(conn)
    return conn

# COMMAND ----------
# MAGIC %md
# MAGIC ## Detection and normalization helpers

# COMMAND ----------
def infer_source_type(path: str) -> str:
    p = path.lower()
    if "/external/" in p or "/osfi/" in p:
        return "external_regulatory"
    if "/policy/" in p:
        return "internal_policy"
    return "internal_reference"


def infer_doc_type(source_type: str) -> str:
    """Map source_type to the metadata spec's DocType ('regulatory' | 'internal')."""
    if source_type == "external_regulatory":
        return "regulatory"
    return "internal"


def contains_requirement(text: str) -> bool:
    t = text.lower()
    return any(k in t for k in [" must ", " should ", " shall ", " required ", " requirement "])


def is_definition_like(text: str) -> bool:
    t = text.strip().lower()
    return t.startswith("definition") or bool(re.match(r"^[A-Za-z\-\s]{1,80}\s+means\s+", t))


def contains_parameter_like(text: str) -> bool:
    t = text.lower()
    return any(k in t for k in ["rate", "%", "bps", "parameter", "threshold", "limit", "minimum", "maximum"])


def contains_deadline_like(text: str) -> bool:
    t = text.lower()
    return any(k in t for k in ["due", "deadline", "submit by", "late", "penalty"])


def contains_assignment_like(text: str) -> bool:
    t = text.lower()
    return any(k in t for k in ["assignment", "project", "deliverable", "submit"])

# COMMAND ----------
def build_registry_row(path: str) -> DocumentRegistryRow:
    raw = read_adls_bytes(path)
    ext = path.split(".")[-1].lower()
    source_type = infer_source_type(path)
    source_system = "osfi_guidance_library" if "osfi" in path.lower() else "internal_repo"

    return DocumentRegistryRow(
        doc_id=os.path.basename(path).replace(".", "_"),
        title=None,
        source_type=source_type,
        raw_path=path,
        raw_file_type=ext,
        source_system=source_system,
        source_url=None,
        sha256=sha256_bytes(raw),
        ingested_at=utc_now_iso(),
    )

# COMMAND ----------
# MAGIC %md
# MAGIC ## Parsers

# COMMAND ----------
def parse_osfi_canonical_json(path: str) -> CanonicalDocument:
    obj = json.loads(read_adls_bytes(path).decode("utf-8"))

    units: list[CanonicalUnit] = []
    for sec_idx, sec in enumerate(obj.get("sections", []), start=1):
        section_heading = (sec.get("heading") or "").strip()
        base_heading_path = [obj.get("title", "")]
        if section_heading:
            base_heading_path.append(section_heading)
        section_path = " > ".join(p for p in base_heading_path if p)

        for item_idx, item in enumerate(sec.get("content", []), start=1):
            if item.get("type") == "paragraph":
                text = (item.get("text") or "").strip()
                if text:
                    unit_id = f"{obj['doc_id']}::sec{sec_idx}::p{item_idx}"
                    units.append(CanonicalUnit(
                        unit_id=unit_id,
                        unit_type="paragraph",
                        heading_path=base_heading_path[:],
                        section_path=section_path,
                        text=text,
                        citation_anchor=unit_id,
                        paragraph_anchor=f"sec{sec_idx}.p{item_idx}",
                        contains_requirement=contains_requirement(f" {text} "),
                        contains_definition=is_definition_like(text),
                        contains_parameter=contains_parameter_like(text),
                    ))
            elif item.get("type") == "list":
                list_text = "\n".join(i.strip() for i in item.get("items", []) if i.strip())
                if list_text:
                    unit_id = f"{obj['doc_id']}::sec{sec_idx}::list{item_idx}"
                    units.append(CanonicalUnit(
                        unit_id=unit_id,
                        unit_type="list",
                        heading_path=base_heading_path[:],
                        section_path=section_path,
                        text=list_text,
                        citation_anchor=unit_id,
                        paragraph_anchor=f"sec{sec_idx}.list{item_idx}",
                        contains_requirement=True,
                    ))

    related = [
        x for x in [obj.get("supersedes_doc_id"), obj.get("superseded_by_doc_id")] if x
    ]

    return CanonicalDocument(
        doc_id=obj["doc_id"],
        doc_type="regulatory",
        title=obj["title"],
        source_type="external_regulatory",
        document_class=obj.get("document_class", "regulatory_guidance"),
        raw_path=obj.get("raw_html_path") or path,
        canonical_path=obj.get("canonical_json_path") or path,
        source_system=obj.get("source_system"),
        source_url=obj.get("canonical_url") or obj.get("detail_url"),
        short_title=obj.get("guideline_number") or obj.get("short_title"),
        version=obj.get("version_id"),
        status=obj.get("status"),
        effective_date_start=obj.get("effective_date_start"),
        effective_date_end=obj.get("effective_date_end"),
        jurisdiction=obj.get("jurisdiction") or "Canada",
        authority_class=obj.get("authority_class"),
        authority_level=obj.get("authority_level"),
        nova_tier=obj.get("nova_tier"),
        confidentiality="public_regulatory",
        approval_status="official",
        business_owner=obj.get("regulator") or "OSFI",
        audience="Treasury, Risk, Finance",
        regulator=obj.get("regulator") or "OSFI",
        regulator_acronym=obj.get("regulator_acronym") or "OSFI",
        doc_family_id=obj.get("doc_family_id"),
        version_id=obj.get("version_id"),
        version_label=obj.get("version_label"),
        version_sort_key=obj.get("version_sort_key"),
        guideline_number=obj.get("guideline_number"),
        current_version_flag=obj.get("current_version_flag"),
        sector=obj.get("sector"),
        supersedes_doc_id=obj.get("supersedes_doc_id"),
        superseded_by_doc_id=obj.get("superseded_by_doc_id"),
        related_doc_ids=related,
        sha256=obj.get("raw_html_sha256"),
        parser_version=obj.get("parser_version"),
        quality_score=obj.get("quality_score"),
        quality_flags=obj.get("quality_flags", []),
        units=units,
    )

# COMMAND ----------
def parse_html(path: str) -> CanonicalDocument:
    html = read_adls_bytes(path)
    soup = BeautifulSoup(html, "html.parser")

    title = soup.title.get_text(" ", strip=True) if soup.title else os.path.basename(path)
    meta = {}
    for tag in soup.find_all("meta"):
        name = tag.get("name") or tag.get("property")
        content = tag.get("content")
        if name and content is not None:
            meta[name] = content

    heading_stack: list[str] = [title]
    units: list[CanonicalUnit] = []
    counter = 0

    for node in soup.find_all(["h1", "h2", "h3", "h4", "p", "li"]):
        text = node.get_text(" ", strip=True)
        if not text:
            continue
        if node.name.startswith("h"):
            level = int(node.name[1])
            heading_stack = heading_stack[:level - 1]
            heading_stack.append(text)
        else:
            counter += 1
            section_path = " > ".join(heading_stack)
            unit_id = f"{os.path.basename(path)}::u{counter}"
            units.append(CanonicalUnit(
                unit_id=unit_id,
                unit_type="paragraph" if node.name == "p" else "list_item",
                heading_path=heading_stack[:],
                section_path=section_path,
                text=text,
                citation_anchor=unit_id,
                paragraph_anchor=f"u{counter}",
                contains_requirement=contains_requirement(f" {text} "),
                contains_definition=is_definition_like(text),
                contains_parameter=contains_parameter_like(text),
            ))

    source_type = infer_source_type(path)
    return CanonicalDocument(
        doc_id=meta.get("osfi-doc-id", os.path.basename(path).replace(".", "_")),
        doc_type=infer_doc_type(source_type),
        title=title,
        source_type=source_type,
        document_class="html_document",
        raw_path=path,
        canonical_path=path,
        source_system="osfi_guidance_library" if "osfi" in path.lower() else "internal_repo",
        source_url=meta.get("canonical-url"),
        short_title=meta.get("guideline-number"),
        status=meta.get("status"),
        effective_date_start=meta.get("effective-date-start"),
        jurisdiction=meta.get("jurisdiction") or "Canada",
        authority_class=meta.get("authority-class"),
        authority_level=int(meta["authority-level"]) if meta.get("authority-level") else None,
        nova_tier=int(meta["nova-tier"]) if meta.get("nova-tier") else None,
        confidentiality="public_regulatory" if "osfi" in path.lower() else "internal",
        approval_status="official" if "osfi" in path.lower() else "unknown",
        business_owner="OSFI" if "osfi" in path.lower() else None,
        regulator="OSFI" if "osfi" in path.lower() else None,
        regulator_acronym="OSFI" if "osfi" in path.lower() else None,
        guideline_number=meta.get("guideline-number"),
        units=units,
    )

# COMMAND ----------
def parse_docx(path: str, enrichment: Optional[dict[str, Any]] = None) -> CanonicalDocument:
    raw = read_adls_bytes(path)
    doc = DocxDocument(io.BytesIO(raw))
    enrichment = enrichment or {}

    units: list[CanonicalUnit] = []
    heading_path: list[str] = []
    counter = 0
    title = enrichment.get("title") or os.path.basename(path)

    for p in doc.paragraphs:
        text = p.text.strip()
        if not text:
            continue
        style_name = (p.style.name or "").lower() if p.style else ""
        if "title" in style_name and not heading_path:
            title = text
            heading_path = [title]
            continue
        if "heading" in style_name:
            m = re.search(r"heading\s*(\d+)", style_name)
            level = int(m.group(1)) if m else 1
            if not heading_path:
                heading_path = [title]
            heading_path = heading_path[:level]
            if heading_path and heading_path[0] != title:
                heading_path = [title] + heading_path[1:]
            if len(heading_path) == 0:
                heading_path = [title]
            if len(heading_path) < level + 1:
                heading_path += [text]
            else:
                heading_path[level] = text
            heading_path = heading_path[: level + 1]
            continue

        counter += 1
        current_heading = heading_path[:] if heading_path else [title]
        section_path = " > ".join(current_heading)
        unit_id = f"{os.path.basename(path)}::u{counter}"
        units.append(CanonicalUnit(
            unit_id=unit_id,
            unit_type="paragraph",
            heading_path=current_heading,
            section_path=section_path,
            text=text,
            citation_anchor=unit_id,
            paragraph_anchor=f"u{counter}",
            contains_deadline=contains_deadline_like(text),
            contains_assignment=contains_assignment_like(text),
            contains_requirement=contains_requirement(f" {text} "),
            contains_definition=is_definition_like(text),
            contains_parameter=contains_parameter_like(text),
        ))

    source_type = enrichment.get("source_type", infer_source_type(path))
    return CanonicalDocument(
        doc_id=enrichment.get("doc_id", os.path.basename(path).replace(".", "_")),
        doc_type=infer_doc_type(source_type),
        title=title,
        source_type=source_type,
        document_class=enrichment.get("document_class", "internal_reference"),
        raw_path=path,
        canonical_path=enrichment.get("canonical_path", path.replace("bronze/", "silver/canonical_docs/") + ".json"),
        source_system=enrichment.get("source_system", "internal_repo"),
        source_url=enrichment.get("source_url"),
        short_title=enrichment.get("short_title"),
        version=enrichment.get("version"),
        status=enrichment.get("status", "active"),
        effective_date_start=enrichment.get("effective_date_start"),
        effective_date_end=enrichment.get("effective_date_end"),
        jurisdiction=enrichment.get("jurisdiction", "Canada"),
        authority_class=enrichment.get("authority_class"),
        authority_level=enrichment.get("authority_level"),
        nova_tier=enrichment.get("nova_tier"),
        confidentiality=enrichment.get("confidentiality", "internal"),
        approval_status=enrichment.get("approval_status", "unknown"),
        business_owner=enrichment.get("business_owner"),
        business_line=enrichment.get("business_line"),
        legal_entity=enrichment.get("legal_entity"),
        audience=enrichment.get("audience"),
        document_owner=enrichment.get("document_owner"),
        related_doc_ids=enrichment.get("related_doc_ids", []),
        sha256=sha256_bytes(raw),
        parser_version="docx-parser-v1",
        quality_score=0.95,
        units=units,
    )

# COMMAND ----------
def parse_pdf_with_document_intelligence(path: str, enrichment: Optional[dict[str, Any]] = None) -> CanonicalDocument:
    raw = read_adls_bytes(path)
    result = get_doc_intelligence_client().begin_analyze_document(
        model_id="prebuilt-layout",
        body=raw,
    ).result()

    enrichment = enrichment or {}
    title = enrichment.get("title") or os.path.basename(path)
    units: list[CanonicalUnit] = []

    paragraphs = getattr(result, "paragraphs", []) or []
    for idx, p in enumerate(paragraphs, start=1):
        text = (p.content or "").strip()
        if not text:
            continue
        unit_id = f"{os.path.basename(path)}::p{idx}"
        units.append(CanonicalUnit(
            unit_id=unit_id,
            unit_type="paragraph",
            heading_path=[title],
            section_path=title,
            text=text,
            citation_anchor=unit_id,
            paragraph_anchor=f"p{idx}",
            contains_requirement=contains_requirement(f" {text} "),
            contains_definition=is_definition_like(text),
            contains_parameter=contains_parameter_like(text),
        ))

    source_type = enrichment.get("source_type", infer_source_type(path))
    return CanonicalDocument(
        doc_id=enrichment.get("doc_id", os.path.basename(path).replace(".", "_")),
        doc_type=infer_doc_type(source_type),
        title=title,
        source_type=source_type,
        document_class=enrichment.get("document_class", "pdf_document"),
        raw_path=path,
        canonical_path=enrichment.get("canonical_path", path.replace("bronze/", "silver/canonical_docs/") + ".json"),
        source_system=enrichment.get("source_system"),
        source_url=enrichment.get("source_url"),
        short_title=enrichment.get("short_title"),
        version=enrichment.get("version"),
        status=enrichment.get("status"),
        effective_date_start=enrichment.get("effective_date_start"),
        effective_date_end=enrichment.get("effective_date_end"),
        jurisdiction=enrichment.get("jurisdiction", "Canada"),
        authority_class=enrichment.get("authority_class"),
        authority_level=enrichment.get("authority_level"),
        nova_tier=enrichment.get("nova_tier"),
        confidentiality=enrichment.get("confidentiality"),
        approval_status=enrichment.get("approval_status"),
        business_owner=enrichment.get("business_owner"),
        audience=enrichment.get("audience"),
        sha256=sha256_bytes(raw),
        parser_version="document-intelligence-layout-v1",
        quality_score=0.9,
        units=units,
    )

# COMMAND ----------
# MAGIC %md
# MAGIC ## Canonical serialization

# COMMAND ----------
def canonical_document_to_dict(doc: CanonicalDocument) -> dict[str, Any]:
    payload = asdict(doc)
    payload["units"] = [asdict(u) for u in doc.units]
    return payload

# COMMAND ----------
# MAGIC %md
# MAGIC ## Chunking and embeddings — spec-driven (Three Rules)

# COMMAND ----------
MAX_CHARS = 1800
OVERLAP = 250


def _get_spec_for_doc(doc: CanonicalDocument) -> Dict[str, List[str]]:
    """Return the pre-computed metadata spec for this document's type."""
    if doc.doc_type == "regulatory":
        return REGULATORY_SPEC
    return INTERNAL_SPEC


def _resolve_field(field_name: str, doc: CanonicalDocument, unit: CanonicalUnit) -> Any:
    """Look up a field value on the unit first, then fall back to the document."""
    # Unit-level fields
    if field_name == "heading_path":
        return unit.heading_path
    if field_name == "section_path":
        return unit.section_path
    if field_name == "citation_anchor":
        return unit.citation_anchor
    if hasattr(unit, field_name):
        val = getattr(unit, field_name)
        if val is not None:
            return val
    # Document-level fields
    if hasattr(doc, field_name):
        return getattr(doc, field_name)
    return None


def semantic_header(doc: CanonicalDocument, unit: CanonicalUnit) -> str:
    """Build the compact semantic header using ONLY Rule 1 (embedded) fields.

    This ensures the embedding model sees metadata that changes semantic meaning
    (e.g. regulator, short_title, heading_path) but NOT fields that are only for
    filtering or prompting.
    """
    spec = _get_spec_for_doc(doc)
    embedded_fields = spec["embedded_fields"]

    pieces: list[str] = []
    for field_name in embedded_fields:
        if field_name == "heading_path":
            hp = unit.heading_path or []
            if hp:
                pieces.append(" > ".join(hp[-3:]))
            continue
        if field_name == "section_path":
            # section_path is redundant with heading_path in the header; skip
            continue
        val = _resolve_field(field_name, doc, unit)
        if val is not None and val != "":
            pieces.append(str(val))
    return f"[{' | '.join(pieces)}]" if pieces else ""


def split_text(text: str, max_chars: int = MAX_CHARS, overlap: int = OVERLAP) -> list[str]:
    text = text.strip()
    if len(text) <= max_chars:
        return [text]

    chunks = []
    start = 0
    while start < len(text):
        end = min(len(text), start + max_chars)
        chunks.append(text[start:end])
        if end == len(text):
            break
        start = max(0, end - overlap)
    return chunks


def build_chunk_docs(doc: CanonicalDocument) -> list[dict[str, Any]]:
    """Build chunk documents with metadata separated by role (Rules 1-3).

    Each chunk dict contains:
      - chunk_text / bm25_text: text with Rule 1 semantic header prepended.
      - All Rule 2 index fields: stored for filtering/boosting in ES and PGVector.
      - prompt_header: dict of Rule 3 fields, used at answer time.
      - Operational fields: raw_path, canonical_path, sha256, parser_version, quality_score.

    The chunk dict is a flat structure suitable for direct upsert into both stores.
    """
    spec = _get_spec_for_doc(doc)
    index_fields = spec["index_fields"]
    operational_fields = spec["operational_fields"]

    rows: list[dict[str, Any]] = []
    for unit in doc.units:
        body = (unit.text or "").strip()
        if not body:
            continue
        header = semantic_header(doc, unit)
        parts = split_text(body)
        for idx, part in enumerate(parts, start=1):
            text = f"{header}\n{part}" if header else part
            anchor = f"{doc.doc_id}::{unit.unit_id}::chunk{idx}"

            row: dict[str, Any] = {
                "chunk_text": text,
                "bm25_text": text,
                "citation_anchor": anchor,
                "unit_id": unit.unit_id,
                "unit_type": unit.unit_type,
            }

            # Rule 2: index fields — resolve from doc or unit
            for field_name in index_fields:
                if field_name not in row:
                    row[field_name] = _resolve_field(field_name, doc, unit)

            # Unit-level boolean flags that belong in the index
            for flag in UNIT_LEVEL_INDEX_FLAGS:
                if hasattr(unit, flag) and flag not in row:
                    row[flag] = getattr(unit, flag)

            # Operational fields
            for field_name in operational_fields:
                val = _resolve_field(field_name, doc, unit)
                if val is not None:
                    row[field_name] = val

            # Ensure core operational fields always present
            row.setdefault("raw_path", doc.raw_path)
            row.setdefault("canonical_path", doc.canonical_path)
            row.setdefault("sha256", doc.sha256)
            row.setdefault("parser_version", doc.parser_version)
            row.setdefault("quality_score", doc.quality_score)

            rows.append(row)
    return rows

# COMMAND ----------
def embed_texts(texts: list[str], model: str = EMBED_MODEL, dimensions: int = EMBED_DIMENSIONS) -> list[list[float]]:
    client = get_openai_client()
    response = client.embeddings.create(model=model, input=texts, dimensions=dimensions)
    return [item.embedding for item in response.data]

# COMMAND ----------
# MAGIC %md
# MAGIC ## Elasticsearch storage

# COMMAND ----------
def elastic_mapping() -> dict[str, Any]:
    return {
        "mappings": {
            "properties": {
                "doc_id": {"type": "keyword"},
                "source_type": {"type": "keyword"},
                "title": {"type": "text", "fields": {"keyword": {"type": "keyword"}}},
                "short_title": {"type": "keyword"},
                "document_class": {"type": "keyword"},
                "chunk_text": {"type": "text"},
                "bm25_text": {"type": "text"},
                "unit_id": {"type": "keyword"},
                "unit_type": {"type": "keyword"},
                "heading_path": {"type": "keyword"},
                "section_path": {"type": "text", "fields": {"keyword": {"type": "keyword"}}},
                "paragraph_anchor": {"type": "keyword"},
                "citation_anchor": {"type": "keyword"},
                "status": {"type": "keyword"},
                "effective_date_start": {"type": "date"},
                "effective_date_end": {"type": "date"},
                "jurisdiction": {"type": "keyword"},
                "authority_class": {"type": "keyword"},
                "authority_level": {"type": "integer"},
                "nova_tier": {"type": "integer"},
                "confidentiality": {"type": "keyword"},
                "approval_status": {"type": "keyword"},
                "business_owner": {"type": "keyword"},
                "business_line": {"type": "keyword"},
                "legal_entity": {"type": "keyword"},
                "audience": {"type": "keyword"},
                "regulator": {"type": "keyword"},
                "regulator_acronym": {"type": "keyword"},
                "doc_family_id": {"type": "keyword"},
                "version_id": {"type": "keyword"},
                "version_label": {"type": "keyword"},
                "guideline_number": {"type": "keyword"},
                "current_version_flag": {"type": "keyword"},
                "sector": {"type": "keyword"},
                "supersedes_doc_id": {"type": "keyword"},
                "superseded_by_doc_id": {"type": "keyword"},
                "contains_definition": {"type": "boolean"},
                "contains_formula": {"type": "boolean"},
                "contains_deadline": {"type": "boolean"},
                "contains_requirement": {"type": "boolean"},
                "contains_parameter": {"type": "boolean"},
                "contains_assignment": {"type": "boolean"},
                "dense_vector": {"type": "dense_vector", "index": True, "similarity": "cosine"},
                "raw_path": {"type": "keyword"},
                "canonical_path": {"type": "keyword"},
                "sha256": {"type": "keyword"},
                "parser_version": {"type": "keyword"},
                "quality_score": {"type": "float"},
            }
        }
    }


def ensure_elastic_index(index_name: str = CHUNK_INDEX_NAME) -> None:
    es = get_elastic_client()
    if not es.indices.exists(index=index_name):
        es.indices.create(index=index_name, body=elastic_mapping())


def upsert_chunks_to_elastic(rows: list[dict[str, Any]], index_name: str = CHUNK_INDEX_NAME) -> None:
    ensure_elastic_index(index_name)
    actions = [{"_index": index_name, "_id": row["citation_anchor"], "_source": row} for row in rows]
    bulk(get_elastic_client(), actions)

# COMMAND ----------
# MAGIC %md
# MAGIC ## PGVector storage

# COMMAND ----------
def ensure_pgvector_schema(table_name: str = "nova_chunks") -> None:
    ddl = f"""
    CREATE EXTENSION IF NOT EXISTS vector;
    CREATE TABLE IF NOT EXISTS {table_name} (
        citation_anchor TEXT PRIMARY KEY,
        doc_id TEXT NOT NULL,
        source_type TEXT,
        title TEXT,
        short_title TEXT,
        document_class TEXT,
        chunk_text TEXT NOT NULL,
        heading_path JSONB,
        section_path TEXT,
        unit_id TEXT,
        unit_type TEXT,
        paragraph_anchor TEXT,
        status TEXT,
        effective_date_start DATE NULL,
        effective_date_end DATE NULL,
        jurisdiction TEXT,
        authority_class TEXT,
        authority_level INT NULL,
        nova_tier INT NULL,
        confidentiality TEXT,
        approval_status TEXT,
        business_owner TEXT,
        business_line TEXT,
        legal_entity TEXT,
        audience TEXT,
        regulator TEXT,
        doc_family_id TEXT,
        version_id TEXT,
        version_label TEXT,
        guideline_number TEXT,
        current_version_flag TEXT,
        sector TEXT,
        supersedes_doc_id TEXT,
        superseded_by_doc_id TEXT,
        contains_definition BOOLEAN,
        contains_formula BOOLEAN,
        contains_deadline BOOLEAN,
        contains_requirement BOOLEAN,
        contains_parameter BOOLEAN,
        contains_assignment BOOLEAN,
        raw_path TEXT,
        canonical_path TEXT,
        sha256 TEXT,
        parser_version TEXT,
        quality_score DOUBLE PRECISION,
        embedding vector({EMBED_DIMENSIONS})
    );
    """
    with get_pg_conn() as conn, conn.cursor() as cur:
        cur.execute(ddl)


def upsert_chunks_to_pgvector(rows: list[dict[str, Any]], table_name: str = "nova_chunks") -> None:
    ensure_pgvector_schema(table_name)

    # Build column lists dynamically from the first row so we stay in sync
    # with whatever fields build_chunk_docs() produces.
    # The embedding column is handled specially (renamed from dense_vector).
    SKIP_KEYS = {"bm25_text", "dense_vector"}
    if not rows:
        return
    sample = rows[0]
    col_names = [k for k in sample.keys() if k not in SKIP_KEYS]
    pg_params = [f"%({k})s" if k != "heading_path" else "%(heading_path_json)s" for k in col_names]
    col_sql = ", ".join(col_names)
    val_sql = ", ".join(pg_params)

    sql = f"""
    INSERT INTO {table_name} ({col_sql}, embedding)
    VALUES ({val_sql}, %(dense_vector)s)
    ON CONFLICT (citation_anchor) DO UPDATE SET
        chunk_text = EXCLUDED.chunk_text,
        status = EXCLUDED.status,
        effective_date_start = EXCLUDED.effective_date_start,
        effective_date_end = EXCLUDED.effective_date_end,
        authority_class = EXCLUDED.authority_class,
        authority_level = EXCLUDED.authority_level,
        nova_tier = EXCLUDED.nova_tier,
        approval_status = EXCLUDED.approval_status,
        quality_score = EXCLUDED.quality_score,
        embedding = EXCLUDED.embedding;
    """

    prepared = []
    for row in rows:
        r = dict(row)
        r["heading_path_json"] = json.dumps(r.get("heading_path", []))
        prepared.append(r)

    with get_pg_conn() as conn, conn.cursor() as cur:
        cur.executemany(sql, prepared)

# COMMAND ----------
# MAGIC %md
# MAGIC ## Query helpers — spec-driven prompt rendering (Rule 3)

# COMMAND ----------
def as_of_filter_clauses(as_of_date: str) -> list[dict[str, Any]]:
    return [
        {"range": {"effective_date_start": {"lte": as_of_date}}},
        {
            "bool": {
                "should": [
                    {"bool": {"must_not": {"exists": {"field": "effective_date_end"}}}},
                    {"range": {"effective_date_end": {"gte": as_of_date}}},
                ],
                "minimum_should_match": 1,
            }
        },
    ]


def pgvector_where_clause(
    as_of_date: Optional[str] = None,
    jurisdiction: Optional[str] = None,
    statuses: Optional[list[str]] = None,
    source_types: Optional[list[str]] = None,
) -> tuple[str, list[Any]]:
    clauses = ["1=1"]
    params: list[Any] = []
    if as_of_date:
        clauses.append("(effective_date_start IS NULL OR effective_date_start <= %s)")
        params.append(as_of_date)
        clauses.append("(effective_date_end IS NULL OR effective_date_end >= %s)")
        params.append(as_of_date)
    if jurisdiction:
        clauses.append("jurisdiction = %s")
        params.append(jurisdiction)
    if statuses:
        clauses.append("status = ANY(%s)")
        params.append(statuses)
    if source_types:
        clauses.append("source_type = ANY(%s)")
        params.append(source_types)
    return " AND ".join(clauses), params


def render_hit_for_prompt(hit: dict[str, Any]) -> str:
    """Render a retrieval hit for LLM prompt context using Rule 3 (prompt) fields.

    Determines the doc_type from the hit's source_type, then uses the metadata
    spec to select which fields to inject into the prompt header.
    """
    source_type = hit.get("source_type", "")
    doc_type = "regulatory" if source_type == "external_regulatory" else "internal"
    spec = metadata_spec.build_spec(doc_type)
    prompt_fields = spec["prompt_injected_fields"]

    header_lines: list[str] = []
    for field_name in prompt_fields:
        val = hit.get(field_name)
        if val is None:
            continue
        header_lines.append(f"{field_name.upper()}: {val}")

    # Always include source title for readability
    if "title" not in prompt_fields:
        title = hit.get("title")
        if title:
            header_lines.insert(0, f"SOURCE: {title}")

    return "\n".join(header_lines) + "\n---\n" + hit.get("chunk_text", "")


def build_prompt_context(hits: list[dict[str, Any]]) -> str:
    return "\n\n".join(render_hit_for_prompt(hit) for hit in hits)
