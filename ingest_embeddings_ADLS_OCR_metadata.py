# Databricks notebook source
# MAGIC %md
# MAGIC # Ingest Embeddings ADLS OCR Metadata
# MAGIC
# MAGIC **NOVA RAG Pipeline — Ingestion, OCR & Embedding**
# MAGIC
# MAGIC This notebook handles the full ingestion pipeline:
# MAGIC
# MAGIC ```
# MAGIC File → OCR/Parse → Chunk → Embed → Vector Store (ES + Azure Search) → Master Store
# MAGIC ```
# MAGIC
# MAGIC ### Pipeline Flow
# MAGIC 1. **File Discovery**: Scan ADLS bronze layer for raw files and scraped JSON
# MAGIC 2. **File Processing**: Parse each format (PDF, DOCX, HTML, XLSX, CSV, images, JSON, TXT)
# MAGIC 3. **OCR**: Azure Document Intelligence (prebuilt-layout) + GPT-4o/GPT-5-mini vision fallback
# MAGIC 4. **Chunking**: Heading-aware text splitting with structural metadata
# MAGIC 5. **Embedding**: OpenAI text-embedding-3-large
# MAGIC 6. **Storage**: Upsert to Elasticsearch + Azure Cognitive Search + master document store
# MAGIC
# MAGIC ### NOVA Three Rules Integration
# MAGIC - **Rule 1 (Embed)**: Semantic header prepended to chunk text before embedding
# MAGIC - **Rule 2 (Index)**: NOVA metadata fields stored as ES/Search index fields for filtering
# MAGIC - **Rule 3 (Prompt)**: Metadata injected at prompt time for LLM reasoning
# MAGIC
# MAGIC ### ADLS Layout (Medallion Architecture)
# MAGIC ```
# MAGIC nova-docs/
# MAGIC ├── bronze/
# MAGIC │   ├── external/
# MAGIC │   │   ├── osfi/json/          ← scraped regulatory metadata JSON
# MAGIC │   │   ├── osfi/raw/           ← original PDFs/HTML
# MAGIC │   │   ├── pra/json/
# MAGIC │   │   └── pra/raw/
# MAGIC │   └── internal/              ← DOCX, PDF, HTML, XLSX raw files
# MAGIC ├── silver/
# MAGIC │   ├── canonical_json/        ← parsed canonical documents
# MAGIC │   └── metadata/             ← extracted metadata JSON
# MAGIC └── gold/
# MAGIC     └── chunks/               ← embedded chunk manifests
# MAGIC ```

# COMMAND ----------
# MAGIC %pip install openai pymupdf elasticsearch azure-ai-formrecognizer azure-identity azure-storage-blob python-docx openpyxl beautifulsoup4 langchain langchain-community tiktoken lxml Pillow --quiet

# COMMAND ----------
# MAGIC %md
# MAGIC ## Imports

# COMMAND ----------
import os
import sys
import json
import time
import hashlib
import logging
import base64
import io
import re
import tempfile
import traceback
from pathlib import Path
from datetime import datetime
from typing import List, Dict, Optional
from dataclasses import dataclass, field
from concurrent.futures import ThreadPoolExecutor, as_completed

import fitz  # PyMuPDF
import openai
from openai import OpenAI, AzureOpenAI
import pandas as pd
import requests
import tiktoken

from elasticsearch import Elasticsearch, helpers

from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.document_loaders import (
    TextLoader, UnstructuredHTMLLoader, PyPDFLoader,
    Docx2txtLoader, CSVLoader, UnstructuredExcelLoader,
)

from azure.storage.blob import BlobServiceClient, ContainerClient
from azure.ai.formrecognizer import DocumentAnalysisClient
from azure.identity import DefaultAzureCredential, ClientSecretCredential
from azure.core.credentials import AzureKeyCredential

from bs4 import BeautifulSoup
from docx import Document as DocxDocument
from PIL import Image

# Add current directory to Python path for imports when running standalone
script_dir = Path(__file__).parent if "__file__" in dir() else Path(".")
if str(script_dir) not in sys.path:
    sys.path.insert(0, str(script_dir))

# Try to import OCR processor (optional dependency)
try:
    from ocr_processor import OCRProcessor
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False

# COMMAND ----------
# MAGIC %md
# MAGIC ## Logging Configuration

# COMMAND ----------
# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s - %(name)s - %(levelname)s - %(message)s",
)

# Override and existing logging configuration
for handler in logging.root.handlers[:]:
    logging.root.removeHandler(handler)

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s - %(name)s - %(levelname)s - %(message)s",
)

# Set root logger to ensure all logs are captured
logging.getLogger().setLevel(logging.INFO)

# Suppress verbose HTTP request logs from Azure and other HTTP libraries
logging.getLogger("azure").setLevel(logging.WARNING)
logging.getLogger("urllib3").setLevel(logging.WARNING)
logging.getLogger("requests").setLevel(logging.WARNING)

# Suppress verbose Azure SDK logs
logging.getLogger("azure.core.pipeline.policies.http_logging_policy").setLevel(logging.WARNING)

# Suppress verbose Elasticsearch logs
logging.getLogger("elastic_transport.transport").setLevel(logging.WARNING)

LOG_FILE_PATH = None

logger = logging.getLogger("nova_ingestion")

def log_and_print(message, level="info"):
    """Log and print a message for visibility in notebooks."""
    getattr(logger, level, logger.info)(message)
    try:
        print(message)
    except Exception:
        pass

# COMMAND ----------
# MAGIC %md
# MAGIC ## Configuration & Credentials

# COMMAND ----------
@dataclass
class ProcessingConfig:
    """Configuration for the ingestion pipeline."""
    # ADLS settings
    adls_account_url: str = ""
    adls_file_system: str = "nova-docs"
    adls_container_name: str = "nova-docs"

    # Azure Document Intelligence
    doc_intel_endpoint: str = ""
    doc_intel_key: str = ""

    # OpenAI / Azure OpenAI
    openai_api_key: str = ""
    openai_endpoint: str = ""
    openai_api_version: str = "2024-06-01"
    embed_model: str = "text-embedding-3-large"
    embed_dimensions: int = 1024
    vision_model: str = "gpt-5-mini-2025-08-07-eastus-dz"

    # Elasticsearch
    es_endpoint: str = ""
    es_password: str = ""
    es_index_name: str = "nova_chunks_v1"
    es_cloud_id: str = ""

    # Azure Cognitive Search
    search_endpoint: str = ""
    search_api_key: str = ""
    search_index_name: str = "nova-search-index"

    # Processing
    chunk_size: int = 1000
    chunk_overlap: int = 200
    max_tokens_per_chunk: int = 8191
    batch_size: int = 100
    parallel_pages: int = 4
    parallel_images: int = 4

    # Proxy settings (for network-restricted environments)
    proxy_url: str = ""
    use_proxy: bool = False

    # Counters
    files_processed: int = 0
    files_errored: int = 0
    total_chunks: int = 0

    # NOVA-specific
    path_category: str = "auto"  # "regulatory_json", "internal_raw", "auto"

    def get_proxy_settings(self) -> dict:
        """Return proxy settings for HTTP clients."""
        if self.use_proxy and self.proxy_url:
            return {"http": self.proxy_url, "https": self.proxy_url}
        return {}


# ---- Read config from Databricks widgets or environment ----
def _widget(name: str, default: str = "") -> str:
    try:
        return dbutils.widgets.get(name) or default
    except Exception:
        return os.environ.get(f"NOVA_{name.upper()}", default)


def _secret(scope: str, key: str, env_fallback: str = "") -> str:
    try:
        return dbutils.secrets.get(scope=scope, key=key)
    except Exception:
        return os.environ.get(env_fallback, "")


def load_config() -> ProcessingConfig:
    """Load configuration from Databricks widgets/secrets or environment."""
    try:
        dbutils.widgets.text("adls_account_url", "")
        dbutils.widgets.text("adls_file_system", "nova-docs")
        dbutils.widgets.text("doc_intel_endpoint", "")
        dbutils.widgets.text("elastic_url", "")
        dbutils.widgets.text("chunk_index_name", "nova_chunks_v1")
        dbutils.widgets.text("embed_model", "text-embedding-3-large")
        dbutils.widgets.text("embed_dimensions", "1024")
        dbutils.widgets.text("vision_model", "gpt-5-mini-2025-08-07-eastus-dz")
        dbutils.widgets.text("input_paths_json", "")
        dbutils.widgets.text("search_endpoint", "")
        dbutils.widgets.text("search_index_name", "nova-search-index")
        dbutils.widgets.text("proxy_url", "")
    except Exception:
        pass

    return ProcessingConfig(
        adls_account_url=_widget("adls_account_url"),
        adls_file_system=_widget("adls_file_system", "nova-docs"),
        doc_intel_endpoint=_widget("doc_intel_endpoint"),
        doc_intel_key=_secret("nova-kv", "doc-intel-key", "DOC_INTEL_KEY"),
        openai_api_key=_secret("nova-kv", "openai-api-key", "OPENAI_API_KEY"),
        openai_endpoint=_widget("openai_endpoint", os.environ.get("AZURE_OPENAI_ENDPOINT", "")),
        embed_model=_widget("embed_model", "text-embedding-3-large"),
        embed_dimensions=int(_widget("embed_dimensions", "1024")),
        vision_model=_widget("vision_model", "gpt-5-mini-2025-08-07-eastus-dz"),
        es_endpoint=_widget("elastic_url"),
        es_password=_secret("nova-kv", "es-password", "ELASTICSEARCH_PASSWORD"),
        es_index_name=_widget("chunk_index_name", "nova_chunks_v1"),
        es_cloud_id=_secret("nova-kv", "es-cloud-id", "ELASTICSEARCH_CLOUD_ID"),
        search_endpoint=_widget("search_endpoint"),
        search_api_key=_secret("nova-kv", "search-api-key", "SEARCH_API_KEY"),
        search_index_name=_widget("search_index_name", "nova-search-index"),
        proxy_url=_widget("proxy_url"),
        use_proxy=bool(_widget("proxy_url")),
    )

config = load_config()

# COMMAND ----------
# MAGIC %md
# MAGIC ## Client Initialization

# COMMAND ----------
# ---- ADLS Client ----
_blob_service_client = None
_container_client = None

def get_blob_service_client():
    global _blob_service_client
    if _blob_service_client is None:
        credential = DefaultAzureCredential()
        _blob_service_client = BlobServiceClient(
            account_url=config.adls_account_url,
            credential=credential,
        )
    return _blob_service_client


def get_container_client():
    global _container_client
    if _container_client is None:
        _container_client = get_blob_service_client().get_container_client(config.adls_container_name)
    return _container_client


def check_file_exists_in_adls(file_path):
    """Check if a file path already exists in the filesystem/index."""
    try:
        blob_client = get_container_client().get_blob_client(file_path)
        blob_client.get_blob_properties()
        return True
    except Exception:
        return False


def read_adls_bytes(file_path: str) -> bytes:
    """Read raw bytes from ADLS."""
    blob_client = get_container_client().get_blob_client(file_path)
    return blob_client.download_blob().readall()


def write_adls_bytes(file_path: str, data: bytes):
    """Write bytes to ADLS."""
    blob_client = get_container_client().get_blob_client(file_path)
    blob_client.upload_blob(data, overwrite=True)


def write_adls_json(file_path: str, obj: dict):
    """Write JSON to ADLS."""
    write_adls_bytes(file_path, json.dumps(obj, indent=2, default=str).encode("utf-8"))


def list_adls_files(prefix: str, extensions: list = None) -> list:
    """List files in ADLS under a prefix, optionally filtered by extension."""
    container = get_container_client()
    blobs = container.list_blobs(name_starts_with=prefix)
    files = []
    for blob in blobs:
        if extensions:
            if any(blob.name.lower().endswith(ext) for ext in extensions):
                files.append(blob.name)
        else:
            files.append(blob.name)
    return files


ADLS_TOTAL_FILES_IN_ANALYZE = 0

# ---- OpenAI Client ----
_openai_client = None

def get_openai_client():
    global _openai_client
    if _openai_client is not None:
        return _openai_client
    if config.openai_endpoint:
        _openai_client = AzureOpenAI(
            azure_endpoint=config.openai_endpoint,
            api_key=config.openai_api_key,
            api_version=config.openai_api_version,
        )
    else:
        _openai_client = OpenAI(api_key=config.openai_api_key)
    return _openai_client


# ---- Document Intelligence Client ----
_doc_intel_client = None

def get_doc_intelligence_client():
    global _doc_intel_client
    if _doc_intel_client is None:
        if config.doc_intel_key:
            credential = AzureKeyCredential(config.doc_intel_key)
        else:
            credential = DefaultAzureCredential()
        _doc_intel_client = DocumentAnalysisClient(
            endpoint=config.doc_intel_endpoint,
            credential=credential,
        )
    return _doc_intel_client


# ---- Elasticsearch Client ----
_es_client = None

def get_es_client():
    """Create Elasticsearch client with proxy support."""
    global _es_client
    if _es_client is not None:
        return _es_client

    es_kwargs = {
        "request_timeout": 60,
        "verify_certs": True,
        "retry_on_timeout": True,
        "max_retries": 3,
    }

    if config.es_cloud_id:
        es_kwargs["cloud_id"] = config.es_cloud_id
        es_kwargs["basic_auth"] = ("elastic", config.es_password)
    else:
        es_kwargs["hosts"] = [config.es_endpoint]
        if config.es_password:
            es_kwargs["basic_auth"] = ("elastic", config.es_password)

    # Proxy support for network-restricted environments
    if config.use_proxy and config.proxy_url:
        es_kwargs["proxy"] = config.proxy_url

    _es_client = Elasticsearch(**es_kwargs)

    if not _es_client.ping():
        raise ConnectionError(f"Cannot reach Elasticsearch")
    log_and_print(f"Connected to Elasticsearch")
    return _es_client


# ---- Azure Cognitive Search Client ----
def get_search_client():
    """Create Azure Cognitive Search client."""
    try:
        from azure.search.documents import SearchClient
        from azure.search.documents.indexes import SearchIndexClient
        return SearchIndexClient(
            endpoint=config.search_endpoint,
            credential=AzureKeyCredential(config.search_api_key),
        )
    except ImportError:
        log_and_print("azure-search-documents not installed. Azure Search disabled.", "warning")
        return None

# COMMAND ----------
# MAGIC %md
# MAGIC ## Elasticsearch Index Mapping — NOVA Fields

# COMMAND ----------
def create_es_index_with_mapping(es_client, index_name):
    """Create Elasticsearch index with NOVA metadata field mapping."""
    if es_client.indices.exists(index=index_name):
        log_and_print(f"Index '{index_name}' already exists")
        return

    mapping = {
        "mappings": {
            "properties": {
                # Core chunk fields
                "doc_id": {"type": "keyword"},
                "chunk_id": {"type": "keyword"},
                "chunk_text": {"type": "text"},
                "bm25_text": {"type": "text"},
                "source_file": {"type": "keyword"},
                "source_path": {"type": "keyword"},
                "file_type": {"type": "keyword"},
                "title": {"type": "text", "fields": {"keyword": {"type": "keyword"}}},
                "short_title": {"type": "keyword"},
                "chunk_index": {"type": "integer"},
                "total_chunks": {"type": "integer"},
                "page_number": {"type": "integer"},

                # Heading/section structure
                "heading_path": {"type": "keyword"},
                "section_path": {"type": "text", "fields": {"keyword": {"type": "keyword"}}},
                "heading_text": {"type": "text"},

                # NOVA document metadata (Rule 2 — Index fields)
                "document_class": {"type": "keyword"},
                "source_type": {"type": "keyword"},
                "status": {"type": "keyword"},
                "effective_date_start": {"type": "date"},
                "effective_date_end": {"type": "date"},
                "jurisdiction": {"type": "keyword"},
                "authority_class": {"type": "keyword"},
                "authority_level": {"type": "integer"},
                "nova_tier": {"type": "integer"},
                "confidentiality": {"type": "keyword"},
                "approval_status": {"type": "keyword"},
                "business_owner": {"type": "keyword"},
                "business_line": {"type": "keyword"},
                "legal_entity": {"type": "keyword"},
                "audience": {"type": "keyword"},
                "regulator": {"type": "keyword"},
                "regulator_acronym": {"type": "keyword"},
                "doc_family_id": {"type": "keyword"},
                "version_id": {"type": "keyword"},
                "version_label": {"type": "keyword"},
                "guideline_number": {"type": "keyword"},
                "current_version_flag": {"type": "keyword"},
                "sector": {"type": "keyword"},
                "supersedes_doc_id": {"type": "keyword"},
                "superseded_by_doc_id": {"type": "keyword"},

                # NOVA structural metadata (unit-level)
                "contains_definition": {"type": "boolean"},
                "contains_formula": {"type": "boolean"},
                "contains_deadline": {"type": "boolean"},
                "contains_requirement": {"type": "boolean"},
                "contains_parameter": {"type": "boolean"},
                "contains_assignment": {"type": "boolean"},
                "structural_level": {"type": "keyword"},
                "section_number": {"type": "keyword"},
                "depth": {"type": "integer"},
                "parent_section_id": {"type": "keyword"},
                "is_appendix": {"type": "boolean"},
                "normative_weight": {"type": "keyword"},
                "paragraph_role": {"type": "keyword"},
                "cross_references": {"type": "keyword"},

                # Embedding vector
                "dense_vector": {"type": "dense_vector", "dims": config.embed_dimensions, "index": True, "similarity": "cosine"},

                # Operational / bookkeeping
                "raw_path": {"type": "keyword"},
                "canonical_path": {"type": "keyword"},
                "sha256": {"type": "keyword"},
                "parser_version": {"type": "keyword"},
                "quality_score": {"type": "float"},
                "ingestion_timestamp": {"type": "date"},

                # Original metadata fields (preserved from original code)
                "document_type": {"type": "keyword"},
                "content_type": {"type": "keyword"},
                "has_tables": {"type": "boolean"},
                "has_images": {"type": "boolean"},
                "extraction_method": {"type": "keyword"},
                "word_count": {"type": "integer"},
                "char_count": {"type": "integer"},
            }
        }
    }

    es_client.indices.create(index=index_name, body=mapping)
    log_and_print(f"Created ES index '{index_name}' with NOVA mapping")


# COMMAND ----------
# MAGIC %md
# MAGIC ## Azure Cognitive Search Index — NOVA Fields

# COMMAND ----------
def create_azure_search_index(config):
    """Create Azure Cognitive Search index with NOVA metadata fields and vector search."""
    try:
        from azure.search.documents.indexes.models import (
            SearchIndex, SearchField, SearchFieldDataType,
            SearchableField, SimpleField, VectorSearch,
            HnswAlgorithmConfiguration, VectorSearchProfile,
            SemanticConfiguration, SemanticSearch, SemanticPrioritizedFields, SemanticField,
        )
    except ImportError:
        log_and_print("azure-search-documents not installed. Skipping Azure Search index creation.", "warning")
        return

    fields = [
        SimpleField(name="id", type=SearchFieldDataType.String, key=True, filterable=True),
        SearchableField(name="chunk_text", type=SearchFieldDataType.String),
        SearchableField(name="title", type=SearchFieldDataType.String, filterable=True),
        SimpleField(name="doc_id", type=SearchFieldDataType.String, filterable=True),
        SimpleField(name="chunk_index", type=SearchFieldDataType.Int32, filterable=True, sortable=True),
        SimpleField(name="source_file", type=SearchFieldDataType.String, filterable=True),
        SimpleField(name="source_path", type=SearchFieldDataType.String, filterable=True),
        SimpleField(name="file_type", type=SearchFieldDataType.String, filterable=True, facetable=True),
        SimpleField(name="page_number", type=SearchFieldDataType.Int32, filterable=True),

        # NOVA metadata fields (Rule 2 — Index)
        SimpleField(name="regulator", type=SearchFieldDataType.String, filterable=True, facetable=True),
        SimpleField(name="jurisdiction", type=SearchFieldDataType.String, filterable=True, facetable=True),
        SimpleField(name="authority_class", type=SearchFieldDataType.String, filterable=True, facetable=True),
        SimpleField(name="document_class", type=SearchFieldDataType.String, filterable=True, facetable=True),
        SimpleField(name="normative_weight", type=SearchFieldDataType.String, filterable=True, facetable=True),
        SimpleField(name="status", type=SearchFieldDataType.String, filterable=True),
        SimpleField(name="nova_tier", type=SearchFieldDataType.Int32, filterable=True),
        SimpleField(name="business_line", type=SearchFieldDataType.String, filterable=True, facetable=True),
        SimpleField(name="guideline_number", type=SearchFieldDataType.String, filterable=True),

        # Content flags
        SimpleField(name="contains_requirement", type=SearchFieldDataType.Boolean, filterable=True),
        SimpleField(name="contains_definition", type=SearchFieldDataType.Boolean, filterable=True),

        # Embedding vector
        SearchField(
            name="dense_vector",
            type=SearchFieldDataType.Collection(SearchFieldDataType.Single),
            searchable=True,
            vector_search_dimensions=config.embed_dimensions,
            vector_search_profile_name="nova-vector-profile",
        ),

        # Bookkeeping
        SimpleField(name="ingestion_timestamp", type=SearchFieldDataType.DateTimeOffset, filterable=True, sortable=True),
        SimpleField(name="quality_score", type=SearchFieldDataType.Double, filterable=True),
    ]

    vector_search = VectorSearch(
        algorithms=[HnswAlgorithmConfiguration(name="nova-hnsw")],
        profiles=[VectorSearchProfile(name="nova-vector-profile", algorithm_configuration_name="nova-hnsw")],
    )

    semantic_config = SemanticConfiguration(
        name="nova-semantic-config",
        prioritized_fields=SemanticPrioritizedFields(
            content_fields=[SemanticField(field_name="chunk_text")],
            title_field=SemanticField(field_name="title"),
        ),
    )

    index = SearchIndex(
        name=config.search_index_name,
        fields=fields,
        vector_search=vector_search,
        semantic_search=SemanticSearch(configurations=[semantic_config]),
    )

    search_client = get_search_client()
    if search_client:
        try:
            search_client.create_or_update_index(index)
            log_and_print(f"Created/updated Azure Search index '{config.search_index_name}'")
        except Exception as e:
            log_and_print(f"Error creating Azure Search index: {e}", "error")

# COMMAND ----------
# MAGIC %md
# MAGIC ## Embedding

# COMMAND ----------
def embed_texts(texts: list, model: str = None, dimensions: int = None) -> list:
    """Embed a list of texts using OpenAI."""
    model = model or config.embed_model
    dimensions = dimensions or config.embed_dimensions
    client = get_openai_client()
    response = client.embeddings.create(model=model, input=texts, dimensions=dimensions)
    return [item.embedding for item in response.data]

# COMMAND ----------
# MAGIC %md
# MAGIC ## OCR & Vision — GPT-4o / GPT-5-mini

# COMMAND ----------
def image_to_base64(image_bytes: bytes, image_format: str = "png") -> str:
    """Convert image bytes to base64 string."""
    return base64.b64encode(image_bytes).decode("utf-8")


def get_gpt4o_description(image_bytes: bytes, image_format: str = "png") -> str:
    """Use GPT-5-mini (or configured vision model) to extract text from images.

    GPT-5-mini provides significantly better vision capabilities:
      - 64k max output tokens (no truncation on dense regulatory pages)
      - Better table extraction and structural preservation
      - Improved multilingual support (English/French regulatory docs)

    Falls back to GPT-4o if configured model is unavailable.
    """
    client = get_openai_client()
    b64_image = image_to_base64(image_bytes)
    mime_type = f"image/{image_format}" if image_format != "jpg" else "image/jpeg"

    response = client.chat.completions.create(
        model=config.vision_model,
        messages=[{
            "role": "user",
            "content": [
                {
                    "type": "text",
                    "text": (
                        "Extract ALL text from this document image. "
                        "Preserve the structure including headings, paragraphs, "
                        "lists, tables, and any other formatted content. "
                        "If this is a chart, graph, or diagram: describe key insights and extract data labels. "
                        "If this is a table: preserve the table structure in markdown. "
                        "Be comprehensive and precise."
                    ),
                },
                {
                    "type": "image_url",
                    "image_url": {"url": f"data:{mime_type};base64,{b64_image}", "detail": "high"},
                },
            ],
        }],
        max_tokens=16384,
    )
    return response.choices[0].message.content or ""


def ocr_with_confidence(image_bytes: bytes, image_format: str = "png") -> tuple:
    """Run vision-model OCR and return (text, confidence_score).

    Confidence heuristic (0.0-1.0):
      - 1.0 = dense, coherent English/French text
      - 0.7+ = normal document text
      - <0.5 = likely garbage / failed OCR (flag for human review)
    """
    text = get_gpt4o_description(image_bytes, image_format)
    if not text or not text.strip():
        return ("", 0.0)

    alpha_chars = sum(1 for c in text if c.isalpha())
    total_chars = len(text)
    alpha_ratio = alpha_chars / max(total_chars, 1)
    words = text.split()
    avg_word_len = sum(len(w) for w in words) / max(len(words), 1)
    ascii_chars = sum(1 for c in text if ord(c) < 128)
    ascii_ratio = ascii_chars / max(total_chars, 1)

    density_score = min(alpha_ratio / 0.65, 1.0)
    wordlen_score = 1.0 if 3.0 <= avg_word_len <= 12.0 else 0.5
    ascii_score = min(ascii_ratio / 0.85, 1.0)
    confidence = round(0.50 * density_score + 0.25 * wordlen_score + 0.25 * ascii_score, 3)

    return (text, confidence)


def extract_text_with_ocr_enhanced(image_bytes: bytes, image_format: str = "png") -> str:
    """Extract text from image using OCR processor if available, else GPT vision."""
    if OCR_AVAILABLE:
        try:
            processor = OCRProcessor()
            result = processor.process(image_bytes)
            if result and len(result.strip()) > 20:
                return result
        except Exception as e:
            log_and_print(f"OCR processor failed, falling back to GPT vision: {e}", "warning")

    return get_gpt4o_description(image_bytes, image_format)

# COMMAND ----------
# MAGIC %md
# MAGIC ## Image Extraction from PDF

# COMMAND ----------
def extract_images_from_pdf(pdf_bytes: bytes, rasterize_if_no_images: bool = True) -> list:
    """Extract images from a PDF using PyMuPDF.

    Two-phase strategy:
      1. Extract embedded image objects via page.get_images() (fast, lossless)
      2. If a page has NO embedded images and rasterize_if_no_images=True,
         render the page to a 300-DPI pixmap (catches scanned PDFs)
    """
    images = []
    try:
        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        for page_num in range(len(doc)):
            page = doc[page_num]
            image_list = page.get_images(full=True)
            page_had_images = False

            for img_idx, img in enumerate(image_list):
                xref = img[0]
                base_image = doc.extract_image(xref)
                if base_image:
                    images.append({
                        "image_bytes": base_image["image"],
                        "page_number": page_num + 1,
                        "image_format": base_image.get("ext", "png"),
                        "source": "embedded_object",
                    })
                    page_had_images = True

            # Phase 2: rasterize the page if no embedded images found
            if not page_had_images and rasterize_if_no_images:
                try:
                    pix = page.get_pixmap(dpi=300)
                    images.append({
                        "image_bytes": pix.tobytes("png"),
                        "page_number": page_num + 1,
                        "image_format": "png",
                        "source": "page_rasterization",
                    })
                except Exception as pix_err:
                    log_and_print(f"  Page {page_num+1} rasterization failed: {pix_err}", "warning")

        doc.close()
    except Exception as e:
        log_and_print(f"Image extraction from PDF failed: {e}", "warning")
    return images


def extract_and_process_images_from_url(html_content, base_url=None):
    """Extract and process images from HTML using GPT-4o/GPT-5-mini OCR.

    Finds image tags in HTML, downloads images, runs OCR on each.
    Returns list of dictionaries with image OCR results.
    """
    soup = BeautifulSoup(html_content, "html.parser")
    image_results = []
    img_tags = soup.find_all("img")

    if not img_tags:
        return image_results

    for idx, img_tag in enumerate(img_tags):
        src = img_tag.get("src", "")
        alt_text = img_tag.get("alt", "")
        if not src:
            continue

        try:
            # Handle data URIs
            if src.startswith("data:"):
                # Extract base64 data from data URI
                header, data = src.split(",", 1)
                image_bytes = base64.b64decode(data)
            else:
                # Build full URL if needed
                if base_url and not src.startswith(("http://", "https://")):
                    full_url = f"{base_url.rstrip('/')}/{src.lstrip('/')}"
                else:
                    full_url = src

                resp = requests.get(full_url, timeout=30)
                if resp.status_code != 200:
                    continue
                image_bytes = resp.content

            # OCR the image
            img_text = get_gpt4o_description(image_bytes, "png")
            if img_text and img_text.strip():
                image_results.append({
                    "image_index": idx,
                    "alt_text": alt_text,
                    "ocr_text": img_text,
                    "source_url": src[:200],
                })
        except Exception as e:
            log_and_print(f"Error processing image {idx}: {e}", "warning")

    return image_results

# COMMAND ----------
# MAGIC %md
# MAGIC ## HTML Processing

# COMMAND ----------
def clean_html_content(html_content: str) -> str:
    """Clean HTML and extract text content using BeautifulSoup."""
    soup = BeautifulSoup(html_content, "html.parser")

    # Remove script and style elements
    for element in soup(["script", "style", "nav", "footer", "header"]):
        element.decompose()

    text = soup.get_text(separator="\n")

    # Clean up whitespace
    lines = [line.strip() for line in text.splitlines()]
    text = "\n".join(line for line in lines if line)
    return text


def extract_tables_from_html(html_content: str) -> list:
    """Extract tables from HTML as markdown."""
    soup = BeautifulSoup(html_content, "html.parser")
    tables = []
    for table in soup.find_all("table"):
        rows = []
        for tr in table.find_all("tr"):
            cells = [td.get_text(strip=True) for td in tr.find_all(["td", "th"])]
            rows.append(" | ".join(cells))
        if rows:
            tables.append("\n".join(rows))
    return tables


def process_html_with_ocr(file_path, file_content, filepath, token_manager, extracted_metadata=None):
    """Process HTML file: extract text, tables, and images with OCR.

    Args:
        file_path: Path to the HTML file
        file_content: Raw HTML content as string
        filepath: Full ADLS path
        token_manager: Token manager for rate limiting
        extracted_metadata: Pre-extracted metadata dict

    Returns:
        List of document objects with extracted content
    """
    extracted_metadata = extracted_metadata or {}

    # Clean and extract text
    text_content = clean_html_content(file_content)

    # Extract tables
    tables = extract_tables_from_html(file_content)
    if tables:
        text_content += "\n\n--- Tables ---\n" + "\n\n".join(tables)

    # Extract and OCR images
    image_results = extract_and_process_images_from_url(file_content)
    if image_results:
        for img_result in image_results:
            text_content += f"\n\n--- Image {img_result['image_index']+1} ---\n"
            if img_result.get("alt_text"):
                text_content += f"Alt: {img_result['alt_text']}\n"
            text_content += img_result.get("ocr_text", "")

    # Build document record
    doc_record = {
        "content": text_content,
        "source_path": filepath,
        "source_file": os.path.basename(filepath),
        "file_type": "html",
        "extraction_method": "beautifulsoup_ocr",
        "has_tables": len(tables) > 0,
        "has_images": len(image_results) > 0,
        "word_count": len(text_content.split()),
        "char_count": len(text_content),
        "metadata": extracted_metadata,
    }

    return [doc_record]

# COMMAND ----------
# MAGIC %md
# MAGIC ## PDF Processing — Document Intelligence + GPT Vision Fallback

# COMMAND ----------
def process_pdf_with_ocr(file_path, file_content, filepath, token_manager, extracted_metadata=None,
                         parallel_pages=None, parallel_images=None):
    """Process PDF using DI to extract text, images, and tables.

    Strategy:
    1. Try Azure Document Intelligence (prebuilt-layout) first
    2. Smart sparse-text detection (chars-per-page ratio)
    3. Extract page images (embedded + rasterized) and run GPT vision OCR
    4. Confidence scoring on OCR output
    5. Merge DI text + OCR results

    Args:
        file_path: Local temp path to the PDF
        file_content: Raw PDF bytes
        filepath: Full ADLS path
        token_manager: Token manager for rate limiting
        extracted_metadata: Pre-extracted metadata dict
        parallel_pages: Number of pages to process in parallel
        parallel_images: Number of images to process in parallel

    Returns:
        List of document objects with extracted content
    """
    parallel_pages = parallel_pages or config.parallel_pages
    parallel_images = parallel_images or config.parallel_images
    extracted_metadata = extracted_metadata or {}

    file_content_bytes = file_content if isinstance(file_content, bytes) else file_content.encode("utf-8")

    # Count pages for smart threshold
    try:
        tmp_doc = fitz.open(stream=file_content_bytes, filetype="pdf")
        page_count = len(tmp_doc)
        tmp_doc.close()
    except Exception:
        page_count = 1

    # Step 1: Try Azure Document Intelligence
    di_text = ""
    di_sections = []
    try:
        log_and_print(f"  Processing PDF with DI: {filepath}")
        client = get_doc_intelligence_client()
        poller = client.begin_analyze_document("prebuilt-layout", document=file_content_bytes)
        result = poller.result()

        # Extract paragraphs with structural info
        paragraphs = getattr(result, "paragraphs", []) or []
        current_heading_path = [os.path.basename(filepath)]

        for idx, p in enumerate(paragraphs):
            text = (p.content or "").strip()
            if not text:
                continue
            role = getattr(p, "role", None)

            section_info = {
                "text": text,
                "role": role,
                "page": getattr(p, "bounding_regions", [{}])[0].page_number if getattr(p, "bounding_regions", None) else None,
                "heading_path": current_heading_path[:],
            }

            if role in ("sectionHeading", "title"):
                if role == "title":
                    current_heading_path = [text]
                else:
                    current_heading_path = current_heading_path[:1] + [text]
                section_info["heading_path"] = current_heading_path[:]
            di_sections.append(section_info)

        di_text = "\n\n".join(s["text"] for s in di_sections)

        # Extract tables from DI result
        tables = getattr(result, "tables", []) or []
        for table in tables:
            table_text = "\n--- Table ---\n"
            for cell in table.cells:
                table_text += f"[R{cell.row_index},C{cell.column_index}] {cell.content}\n"
            di_text += "\n" + table_text

        total_text_len = len(di_text)
        chars_per_page = total_text_len / max(page_count, 1)

        # Smart threshold: at least 100 chars per page on average
        if chars_per_page >= 100:
            log_and_print(f"  DI extracted {total_text_len} chars ({chars_per_page:.0f}/page). Using DI result.")
        else:
            log_and_print(f"  DI sparse text ({total_text_len} chars, {chars_per_page:.0f}/page). Adding GPT vision OCR...")

    except Exception as e:
        log_and_print(f"  DI failed: {e}. Falling back to GPT vision OCR...", "warning")

    # Step 2: OCR for scanned pages (if DI was sparse or failed)
    ocr_text_parts = []
    low_confidence_pages = []
    chars_per_page = len(di_text) / max(page_count, 1) if di_text else 0

    if chars_per_page < 100:
        images = extract_images_from_pdf(file_content_bytes, rasterize_if_no_images=True)
        if images:
            log_and_print(f"  Running GPT vision OCR on {len(images)} images...")
            for img_info in images:
                try:
                    ocr_text, confidence = ocr_with_confidence(
                        img_info["image_bytes"], img_info["image_format"]
                    )
                    if ocr_text.strip():
                        ocr_text_parts.append({
                            "text": ocr_text,
                            "page": img_info["page_number"],
                            "confidence": confidence,
                            "source": img_info.get("source", "unknown"),
                        })
                        if confidence < 0.5:
                            low_confidence_pages.append(img_info["page_number"])
                except Exception:
                    continue

    if low_confidence_pages:
        log_and_print(f"  Warning: Low OCR confidence on pages {low_confidence_pages}. Flag for human review.", "warning")

    # Merge DI + OCR text
    full_text = di_text
    if ocr_text_parts:
        full_text += "\n\n--- OCR Extracted Content ---\n"
        for part in ocr_text_parts:
            full_text += f"\n[Page {part['page']}]\n{part['text']}\n"

    if not full_text.strip():
        log_and_print(f"  No text extracted from PDF: {filepath}", "warning")
        return []

    # Build document record
    doc_record = {
        "content": full_text,
        "source_path": filepath,
        "source_file": os.path.basename(filepath),
        "file_type": "pdf",
        "extraction_method": "document_intelligence_ocr_fallback",
        "has_tables": bool(getattr(result, "tables", None)) if "result" in dir() else False,
        "has_images": len(ocr_text_parts) > 0,
        "word_count": len(full_text.split()),
        "char_count": len(full_text),
        "page_count": page_count,
        "di_sections": di_sections,
        "metadata": extracted_metadata,
        "quality_score": sum(p.get("confidence", 0.9) for p in ocr_text_parts) / max(len(ocr_text_parts), 1) if ocr_text_parts else 0.9,
    }

    return [doc_record]

# COMMAND ----------
# MAGIC %md
# MAGIC ## DOCX Processing

# COMMAND ----------
def process_docx_with_metadata(file_path, file_content, filepath, token_manager, extracted_metadata=None):
    """Process Word document: extract text, tables, metadata, and images.

    Extracts:
    - Full text content with heading structure
    - Tables as markdown
    - Document properties (core.xml metadata)
    - Embedded images via GPT vision OCR
    """
    extracted_metadata = extracted_metadata or {}
    file_content_bytes = file_content if isinstance(file_content, bytes) else file_content.encode("utf-8")

    doc = DocxDocument(io.BytesIO(file_content_bytes))

    # Extract document properties
    core_props = doc.core_properties
    native_metadata = {
        "title": core_props.title or "",
        "author": core_props.author or "",
        "subject": core_props.subject or "",
        "created": str(core_props.created) if core_props.created else "",
        "modified": str(core_props.modified) if core_props.modified else "",
        "last_modified_by": core_props.last_modified_by or "",
        "revision": str(core_props.revision) if core_props.revision else "",
        "category": core_props.category or "",
        "keywords": core_props.keywords or "",
    }

    # Extract text with heading tracking
    text_parts = []
    current_heading = ""
    heading_path = [os.path.basename(filepath)]

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        style_name = para.style.name.lower() if para.style else ""

        if "heading" in style_name:
            current_heading = text
            heading_path = heading_path[:1] + [text]
            text_parts.append(f"\n## {text}\n")
        else:
            text_parts.append(text)

    # Extract tables
    table_texts = []
    for table_idx, table in enumerate(doc.tables):
        table_md = f"\n--- Table {table_idx + 1} ---\n"
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            table_md += " | ".join(cells) + "\n"
        table_texts.append(table_md)

    full_text = "\n".join(text_parts)
    if table_texts:
        full_text += "\n" + "\n".join(table_texts)

    # Extract and OCR embedded images
    image_texts = []
    try:
        # Use fitz to extract images from DOCX (via PDF conversion) or directly
        for rel in doc.part.rels.values():
            if "image" in rel.reltype:
                try:
                    img_data = rel.target_part.blob
                    img_text = get_gpt4o_description(img_data, "png")
                    if img_text and img_text.strip():
                        image_texts.append(img_text)
                except Exception:
                    continue
    except Exception as e:
        log_and_print(f"  DOCX image extraction failed: {e}", "warning")

    if image_texts:
        full_text += "\n\n--- Embedded Images ---\n"
        for i, txt in enumerate(image_texts):
            full_text += f"\n[Image {i+1}]\n{txt}\n"

    doc_record = {
        "content": full_text,
        "source_path": filepath,
        "source_file": os.path.basename(filepath),
        "file_type": "docx",
        "extraction_method": "python-docx",
        "has_tables": len(table_texts) > 0,
        "has_images": len(image_texts) > 0,
        "word_count": len(full_text.split()),
        "char_count": len(full_text),
        "native_metadata": native_metadata,
        "metadata": extracted_metadata,
    }

    return [doc_record]

# COMMAND ----------
# MAGIC %md
# MAGIC ## Excel / CSV Processing

# COMMAND ----------
def process_excel_with_metadata(file_path, file_content, filepath, token_manager, extracted_metadata=None):
    """Process Excel/CSV file: extract sheets as text tables."""
    extracted_metadata = extracted_metadata or {}
    file_content_bytes = file_content if isinstance(file_content, bytes) else file_content.encode("utf-8")
    ext = os.path.splitext(filepath)[1].lower()

    text_parts = []

    try:
        if ext in (".csv", ".tsv"):
            sep = "\t" if ext == ".tsv" else ","
            df = pd.read_csv(io.BytesIO(file_content_bytes), sep=sep)
            text_parts.append(f"Sheet: data\n{df.to_markdown(index=False)}")
        else:
            xls = pd.ExcelFile(io.BytesIO(file_content_bytes))
            for sheet_name in xls.sheet_names:
                df = pd.read_excel(xls, sheet_name=sheet_name)
                text_parts.append(f"Sheet: {sheet_name}\n{df.to_markdown(index=False)}")
    except Exception as e:
        log_and_print(f"  Excel/CSV parsing failed: {e}", "warning")
        return []

    full_text = "\n\n".join(text_parts)

    doc_record = {
        "content": full_text,
        "source_path": filepath,
        "source_file": os.path.basename(filepath),
        "file_type": ext.lstrip("."),
        "extraction_method": "pandas",
        "has_tables": True,
        "has_images": False,
        "word_count": len(full_text.split()),
        "char_count": len(full_text),
        "metadata": extracted_metadata,
    }

    return [doc_record]

# COMMAND ----------
# MAGIC %md
# MAGIC ## JSON / TXT / Markdown Processing

# COMMAND ----------
def process_json_document(file_path, file_content, filepath, token_manager, extracted_metadata=None):
    """Process JSON file: extract content and metadata."""
    extracted_metadata = extracted_metadata or {}

    try:
        if isinstance(file_content, bytes):
            data = json.loads(file_content.decode("utf-8"))
        else:
            data = json.loads(file_content)
    except json.JSONDecodeError as e:
        log_and_print(f"  JSON parse error: {e}", "warning")
        return []

    # Extract text content from JSON
    text = json.dumps(data, indent=2, default=str)

    doc_record = {
        "content": text,
        "source_path": filepath,
        "source_file": os.path.basename(filepath),
        "file_type": "json",
        "extraction_method": "json_parse",
        "has_tables": False,
        "has_images": False,
        "word_count": len(text.split()),
        "char_count": len(text),
        "metadata": extracted_metadata,
        "json_data": data,
    }

    return [doc_record]


def process_text_file(file_path, file_content, filepath, token_manager, extracted_metadata=None):
    """Process plain text or markdown file."""
    extracted_metadata = extracted_metadata or {}

    if isinstance(file_content, bytes):
        text = file_content.decode("utf-8", errors="replace")
    else:
        text = file_content

    ext = os.path.splitext(filepath)[1].lower()

    doc_record = {
        "content": text,
        "source_path": filepath,
        "source_file": os.path.basename(filepath),
        "file_type": ext.lstrip(".") or "txt",
        "extraction_method": "text_read",
        "has_tables": False,
        "has_images": False,
        "word_count": len(text.split()),
        "char_count": len(text),
        "metadata": extracted_metadata,
    }

    return [doc_record]


def process_image_file(file_path, file_content, filepath, token_manager, extracted_metadata=None):
    """Process standalone image file with GPT vision OCR."""
    extracted_metadata = extracted_metadata or {}
    file_content_bytes = file_content if isinstance(file_content, bytes) else file_content.encode("utf-8")
    ext = os.path.splitext(filepath)[1].lower().lstrip(".")

    text = extract_text_with_ocr_enhanced(file_content_bytes, ext or "png")

    doc_record = {
        "content": text,
        "source_path": filepath,
        "source_file": os.path.basename(filepath),
        "file_type": ext or "image",
        "extraction_method": "gpt_vision_ocr",
        "has_tables": False,
        "has_images": True,
        "word_count": len(text.split()),
        "char_count": len(text),
        "metadata": extracted_metadata,
    }

    return [doc_record]

# COMMAND ----------
# MAGIC %md
# MAGIC ## Regulatory Scraped JSON Processing

# COMMAND ----------
def process_regulatory_scraped_json(file_path, file_content, filepath, token_manager, extracted_metadata=None):
    """Process scraped regulatory JSON files (OSFI, PRA, etc.).

    These JSON files already contain rich metadata and parsed content
    from the scraping pipeline. We read the content and map the metadata
    to NOVA fields rather than re-extracting.
    """
    extracted_metadata = extracted_metadata or {}

    try:
        if isinstance(file_content, bytes):
            data = json.loads(file_content.decode("utf-8"))
        else:
            data = json.loads(file_content)
    except json.JSONDecodeError as e:
        log_and_print(f"  JSON parse error in regulatory file: {e}", "warning")
        return []

    # Extract content from scraped JSON structure
    sections = data.get("sections", [])
    text_parts = []

    if sections:
        for section in sections:
            section_title = section.get("title", section.get("heading", ""))
            if section_title:
                text_parts.append(f"\n## {section_title}\n")
            # Handle content items within sections
            content_items = section.get("content", section.get("items", []))
            if isinstance(content_items, list):
                for item in content_items:
                    if isinstance(item, dict):
                        text_parts.append(item.get("text", item.get("content", str(item))))
                    elif isinstance(item, str):
                        text_parts.append(item)
            elif isinstance(content_items, str):
                text_parts.append(content_items)

    # Fallback: try common top-level text fields
    if not text_parts:
        for key in ("content", "text", "body", "description", "summary"):
            if key in data and isinstance(data[key], str) and data[key].strip():
                text_parts.append(data[key])
                break

    if not text_parts:
        text_parts.append(json.dumps(data, indent=2, default=str))

    full_text = "\n".join(text_parts)

    # Map scraped JSON fields to NOVA metadata
    nova_metadata = {
        "doc_id": data.get("doc_id", data.get("id", os.path.basename(filepath).replace(".", "_"))),
        "title": data.get("title", data.get("name", "")),
        "short_title": data.get("short_title", data.get("abbreviation", "")),
        "regulator": data.get("regulator", data.get("authority", _infer_regulator_from_path(filepath))),
        "regulator_acronym": data.get("regulator_acronym", ""),
        "guideline_number": data.get("guideline_number", data.get("reference_number", "")),
        "jurisdiction": data.get("jurisdiction", ""),
        "authority_class": data.get("authority_class", ""),
        "nova_tier": data.get("nova_tier"),
        "status": data.get("status", "active"),
        "effective_date_start": data.get("effective_date_start", data.get("effective_date", data.get("publication_date", ""))),
        "effective_date_end": data.get("effective_date_end", ""),
        "document_class": data.get("document_class", data.get("category", "")),
        "version_id": data.get("version_id", data.get("version", "")),
        "version_label": data.get("version_label", ""),
        "current_version_flag": data.get("current_version_flag", "current"),
        "sector": data.get("sector", ""),
        "supersedes_doc_id": data.get("supersedes_doc_id", data.get("supersedes", "")),
        "superseded_by_doc_id": data.get("superseded_by_doc_id", data.get("superseded_by", "")),
        "doc_family_id": data.get("doc_family_id", ""),
        "source_type": "regulatory",
    }
    # Merge with any pre-extracted metadata
    nova_metadata.update(extracted_metadata)

    doc_record = {
        "content": full_text,
        "source_path": filepath,
        "source_file": os.path.basename(filepath),
        "file_type": "json",
        "extraction_method": "regulatory_scraped_json",
        "has_tables": False,
        "has_images": False,
        "word_count": len(full_text.split()),
        "char_count": len(full_text),
        "metadata": nova_metadata,
        "json_data": data,
    }

    return [doc_record]


def _infer_regulator_from_path(filepath: str) -> str:
    """Infer regulator name from ADLS path."""
    path_lower = filepath.lower()
    if "osfi" in path_lower:
        return "OSFI"
    elif "pra" in path_lower:
        return "PRA"
    elif "boe" in path_lower or "bank_of_england" in path_lower:
        return "BoE"
    elif "bis" in path_lower:
        return "BIS"
    elif "bcbs" in path_lower or "basel" in path_lower:
        return "BCBS"
    return ""

# COMMAND ----------
# MAGIC %md
# MAGIC ## Metadata Extraction from JSON Files

# COMMAND ----------
def extract_metadata_from_json_content_enhanced(json_data: dict) -> dict:
    """Extract NOVA metadata fields from a JSON metadata file.

    Maps common JSON metadata structures to NOVA field names.
    Works with both scraped regulatory JSON and internal metadata JSON.
    """
    metadata = {}

    # Direct field mappings
    field_mappings = {
        "doc_id": ["doc_id", "id", "document_id"],
        "title": ["title", "name", "document_title"],
        "short_title": ["short_title", "abbreviation"],
        "document_class": ["document_class", "category", "type", "document_type"],
        "source_type": ["source_type", "source"],
        "regulator": ["regulator", "authority", "issuing_body"],
        "regulator_acronym": ["regulator_acronym"],
        "guideline_number": ["guideline_number", "reference_number", "ref_number"],
        "jurisdiction": ["jurisdiction", "country"],
        "authority_class": ["authority_class", "authority_type"],
        "nova_tier": ["nova_tier", "tier"],
        "status": ["status", "document_status"],
        "effective_date_start": ["effective_date_start", "effective_date", "publication_date", "date_published"],
        "effective_date_end": ["effective_date_end", "expiry_date"],
        "version_id": ["version_id", "version"],
        "version_label": ["version_label"],
        "business_owner": ["business_owner", "owner", "author"],
        "business_line": ["business_line", "department"],
        "audience": ["audience", "target_audience"],
        "confidentiality": ["confidentiality", "classification"],
        "sector": ["sector", "industry"],
        "doc_family_id": ["doc_family_id", "family_id"],
    }

    for nova_field, source_keys in field_mappings.items():
        for key in source_keys:
            if key in json_data and json_data[key]:
                metadata[nova_field] = json_data[key]
                break

    return metadata


def test_metadata_extraction_from_json_content(json_content):
    """Test metadata extraction from JSON content (validation helper)."""
    try:
        if isinstance(json_content, str):
            data = json.loads(json_content)
        elif isinstance(json_content, bytes):
            data = json.loads(json_content.decode("utf-8"))
        else:
            data = json_content

        metadata = extract_metadata_from_json_content_enhanced(data)
        log_and_print(f"  Extracted {len(metadata)} metadata fields from JSON")
        return metadata
    except Exception as e:
        log_and_print(f"  Metadata extraction test failed: {e}", "warning")
        return {}

# COMMAND ----------
# MAGIC %md
# MAGIC ## Chunking — Heading-Aware Text Splitting

# COMMAND ----------
MIN_CHUNK_SIZE = 100
MAX_CHUNK_SIZE = 2000

def process_chunks(text_content: str, source_file: str = "", metadata: dict = None) -> list:
    """Split text into heading-aware chunks with structural metadata.

    Preserves heading hierarchy and creates chunks that respect
    section boundaries. Each chunk gets heading_path, section_path,
    and content classification metadata.

    Args:
        text_content: Full document text
        source_file: Source filename for chunk IDs
        metadata: Document-level metadata to merge into chunks

    Returns:
        List of chunk dicts with text and metadata
    """
    metadata = metadata or {}
    chunks = []

    if not text_content or not text_content.strip():
        return chunks

    lines = text_content.split("\n")
    current_chunk = ""
    current_heading = ""
    current_headings = []
    heading_path = [source_file] if source_file else []
    chunk_index = 0

    for line in lines:
        stripped = line.strip()
        if not stripped:
            current_chunk += "\n"
            continue

        # Detect headings (markdown-style or ALL CAPS)
        is_heading = False
        if stripped.startswith("#"):
            is_heading = True
            # Count heading level
            level = len(stripped) - len(stripped.lstrip("#"))
            heading_text = stripped.lstrip("# ").strip()
            current_heading = heading_text

            # Update heading path
            if level <= 1:
                heading_path = [source_file, heading_text] if source_file else [heading_text]
            elif level == 2:
                heading_path = heading_path[:2] + [heading_text] if len(heading_path) >= 2 else heading_path + [heading_text]
            else:
                heading_path = heading_path[:min(level, len(heading_path))] + [heading_text]

            current_headings = heading_path[:]

        elif stripped.isupper() and len(stripped) > 3 and len(stripped.split()) <= 10:
            # ALL CAPS likely heading
            is_heading = True
            current_heading = stripped
            heading_path = heading_path[:1] + [stripped]
            current_headings = heading_path[:]

        # If heading starts a new section, flush current chunk
        if is_heading and current_chunk.strip() and len(current_chunk.strip()) >= MIN_CHUNK_SIZE:
            chunk_record = _build_chunk_record(
                current_chunk.strip(), chunk_index, source_file,
                current_headings, current_heading, metadata,
            )
            chunks.append(chunk_record)
            chunk_index += 1
            current_chunk = ""

        current_chunk += line + "\n"

        # If chunk exceeds max size, split it
        if len(current_chunk) >= MAX_CHUNK_SIZE:
            chunk_record = _build_chunk_record(
                current_chunk.strip(), chunk_index, source_file,
                current_headings, current_heading, metadata,
            )
            chunks.append(chunk_record)
            chunk_index += 1
            current_chunk = ""

    # Flush final chunk
    if current_chunk.strip() and len(current_chunk.strip()) >= MIN_CHUNK_SIZE:
        chunk_record = _build_chunk_record(
            current_chunk.strip(), chunk_index, source_file,
            current_headings, current_heading, metadata,
        )
        chunks.append(chunk_record)

    return chunks


def _build_chunk_record(text: str, chunk_index: int, source_file: str,
                        heading_path: list, heading_text: str, metadata: dict) -> dict:
    """Build a chunk record with structural and NOVA metadata."""
    doc_id = metadata.get("doc_id", source_file.replace(".", "_") if source_file else f"doc_{chunk_index}")

    chunk_record = {
        "chunk_id": f"{doc_id}::chunk{chunk_index}",
        "doc_id": doc_id,
        "chunk_text": text,
        "bm25_text": text,
        "chunk_index": chunk_index,
        "source_file": source_file,
        "heading_path": heading_path,
        "section_path": " > ".join(heading_path) if heading_path else "",
        "heading_text": heading_text,

        # Content classification
        "contains_requirement": _contains_requirement(text),
        "contains_definition": _is_definition_like(text),
        "contains_formula": bool(re.search(r"[=×÷∑∫∂]|formula|equation|calculate", text, re.IGNORECASE)),
        "contains_deadline": bool(re.search(r"deadline|due date|by \w+ \d{4}|no later than", text, re.IGNORECASE)),
        "contains_parameter": bool(re.search(r"\b\d+\.?\d*\s*%|\bparameter\b|\bthreshold\b|\blimit\b", text, re.IGNORECASE)),
        "contains_assignment": bool(re.search(r"\bresponsib\w+\b|\bassign\w+\b|\baccountabl\w+\b", text, re.IGNORECASE)),
        "is_appendix": bool(re.search(r"\bappendix\b|\bannex\b|\bschedule\b", " ".join(heading_path), re.IGNORECASE)),

        # Normative weight classification
        "normative_weight": _classify_normative_weight(text),
        "paragraph_role": _classify_paragraph_role(text, heading_text),

        # Cross-references
        "cross_references": _extract_cross_references(text),

        # Word/char counts
        "word_count": len(text.split()),
        "char_count": len(text),
    }

    # Merge document-level metadata
    for key, val in metadata.items():
        if key not in chunk_record and val is not None:
            chunk_record[key] = val

    return chunk_record


def _contains_requirement(text: str) -> bool:
    """Check if text contains regulatory requirement language."""
    patterns = [
        r"\bmust\b", r"\bshall\b", r"\brequired\s+to\b", r"\bis\s+required\b",
        r"\bare\s+required\b", r"\bobligation\b", r"\bmandatory\b",
        r"\bshall\s+not\b", r"\bmust\s+not\b", r"\bprohibited\b",
    ]
    return any(re.search(p, text, re.IGNORECASE) for p in patterns)


def _is_definition_like(text: str) -> bool:
    """Check if text contains definitions."""
    patterns = [
        r'"[^"]+"\s+(means|refers to|is defined as)',
        r"\bmeans\b.*\b(the|a|an)\b",
        r"\bdefined\s+as\b",
        r"\bfor\s+the\s+purpose\s+of\b",
        r"\bdefinition\b",
    ]
    return any(re.search(p, text, re.IGNORECASE) for p in patterns)


def _classify_normative_weight(text: str) -> str:
    """Classify text by normative weight."""
    text_lower = text.lower()
    if any(w in text_lower for w in ("must", "shall", "required", "mandatory", "prohibited")):
        return "mandatory"
    if any(w in text_lower for w in ("should", "expected", "recommended", "encouraged")):
        return "advisory"
    if any(w in text_lower for w in ("may", "can", "permitted", "optional", "at the discretion")):
        return "permissive"
    return "informational"


def _classify_paragraph_role(text: str, heading: str = "") -> str:
    """Classify paragraph role."""
    if _is_definition_like(text):
        return "definition"
    if _contains_requirement(text):
        return "requirement"
    if re.search(r"^\d+\.\s|^[a-z]\)\s|^•\s|^-\s", text):
        return "procedure_step"
    if re.search(r"\bexample\b|\bfor instance\b|\be\.g\.\b|\billustrat", text, re.IGNORECASE):
        return "example"
    if re.search(r"\bexcept\b|\bunless\b|\bnotwithstanding\b|\bexclud", text, re.IGNORECASE):
        return "exception"
    return "narrative"


def _extract_cross_references(text: str) -> list:
    """Extract cross-references from text."""
    refs = []
    # Guideline references (e.g., "B-20", "E-13", "IFRS 9")
    for match in re.finditer(r"\b([A-Z]-\d+|[A-Z]{2,5}\s*\d+)\b", text):
        refs.append(match.group(0))
    # Section references
    for match in re.finditer(r"(?:section|chapter|paragraph|article)\s+(\d+[\.\d]*)", text, re.IGNORECASE):
        refs.append(f"section_{match.group(1)}")
    return list(set(refs))[:20]

# COMMAND ----------
# MAGIC %md
# MAGIC ## NOVA Semantic Header (Rule 1 — Embed)

# COMMAND ----------
def build_semantic_header(chunk: dict) -> str:
    """Build the compact semantic header using Rule 1 (embedded) fields.

    This header is prepended to chunk_text before embedding so the
    vector captures the metadata context.

    Example: [OSFI | B-20 | Residential Mortgage Underwriting | mandatory]
    """
    pieces = []

    # Regulator or business owner
    reg = chunk.get("regulator") or chunk.get("business_owner")
    if reg:
        pieces.append(str(reg))

    # Short title or guideline number
    short = chunk.get("short_title") or chunk.get("guideline_number")
    if short:
        pieces.append(str(short))

    # Document class
    dc = chunk.get("document_class")
    if dc:
        pieces.append(str(dc))

    # Heading path (last 2 levels)
    hp = chunk.get("heading_path", [])
    if hp:
        pieces.append(" > ".join(hp[-2:]))

    # Normative weight
    nw = chunk.get("normative_weight")
    if nw and nw != "informational":
        pieces.append(nw)

    return f"[{' | '.join(pieces)}]" if pieces else ""

# COMMAND ----------
# MAGIC %md
# MAGIC ## Batch Indexing — Elasticsearch & Azure Search

# COMMAND ----------
def index_chunks_and_document(doc_record: dict, chunks: list, es_client=None, index_name: str = None):
    """Index document chunks into Elasticsearch with NOVA metadata.

    Embeds chunk text (with semantic header), then upserts to ES.
    Handles batching and retry logic.
    """
    es_client = es_client or get_es_client()
    index_name = index_name or config.es_index_name

    if not chunks:
        return

    # Build texts for embedding (with semantic header)
    texts_to_embed = []
    for chunk in chunks:
        header = build_semantic_header(chunk)
        text = chunk.get("chunk_text", "")
        embed_text = f"{header}\n{text}" if header else text
        texts_to_embed.append(embed_text)

        # Store the header-enriched text as bm25_text too
        chunk["bm25_text"] = embed_text

    # Embed in batches
    all_embeddings = []
    batch_size = config.batch_size
    for i in range(0, len(texts_to_embed), batch_size):
        batch = texts_to_embed[i:i + batch_size]
        try:
            embeddings = embed_texts(batch)
            all_embeddings.extend(embeddings)
        except Exception as e:
            log_and_print(f"  Embedding batch {i//batch_size} failed: {e}", "error")
            # Fill with None for failed batches
            all_embeddings.extend([None] * len(batch))

    # Build ES actions
    actions = []
    for idx, chunk in enumerate(chunks):
        if idx < len(all_embeddings) and all_embeddings[idx] is not None:
            chunk["dense_vector"] = all_embeddings[idx]
        chunk["ingestion_timestamp"] = datetime.utcnow().isoformat()
        chunk["total_chunks"] = len(chunks)

        action = {
            "_index": index_name,
            "_id": chunk.get("chunk_id", f"chunk_{idx}"),
            "_source": chunk,
        }
        actions.append(action)

    # Upsert with retry
    for attempt in range(1, 4):
        try:
            helpers.bulk(es_client, actions, raise_on_error=True)
            log_and_print(f"  Successfully indexed {len(actions)} chunks to '{index_name}'")
            return
        except Exception as e:
            if attempt == 3:
                log_and_print(f"  ES bulk indexing failed after 3 attempts: {e}", "error")
                raise
            log_and_print(f"  ES bulk attempt {attempt}/3 failed: {e}. Retrying...", "warning")
            time.sleep(2 ** attempt)


def clean_batch(batch_data: list, es_client=None, index_name: str = None) -> list:
    """Clean batch data: filter by source, remove duplicates, validate.

    Batch filtering: check_data["status"] == 'valid', check_data["source"] != 'binary',
    check_data["total_size"] != 0
    """
    cleaned = []
    seen_ids = set()

    for item in batch_data:
        chunk_id = item.get("chunk_id", item.get("_id", ""))

        # Skip duplicates
        if chunk_id in seen_ids:
            continue
        seen_ids.add(chunk_id)

        # Skip empty content
        content = item.get("chunk_text", item.get("content", ""))
        if not content or not content.strip():
            continue

        # Skip very short chunks
        if len(content.strip()) < MIN_CHUNK_SIZE:
            continue

        cleaned.append(item)

    return cleaned

# COMMAND ----------
# MAGIC %md
# MAGIC ## File Discovery — ADLS Bronze Layer

# COMMAND ----------
SUPPORTED_EXTENSIONS = [
    ".pdf", ".docx", ".doc", ".pptx", ".xlsx", ".xls",
    ".csv", ".tsv", ".html", ".htm", ".md", ".txt",
    ".json", ".png", ".jpg", ".jpeg", ".tiff", ".tif", ".bmp", ".gif",
]

def discover_paths_to_ingest(config) -> list:
    """Discover files in ADLS bronze layer for ingestion.

    Returns list of (file_path, path_category) tuples:
      - ("regulatory_json", path) for scraped metadata JSON
      - ("internal_raw", path) for internal raw files
      - ("auto", path) for files that need auto-detection
    """
    tagged_inputs = []

    # 1. External regulatory scraped JSON
    for regulator_prefix in ["bronze/external/osfi/json/", "bronze/external/pra/json/",
                              "bronze/external/boe/json/", "bronze/external/bis/json/"]:
        json_files = list_adls_files(regulator_prefix, extensions=[".json"])
        for f in json_files:
            tagged_inputs.append((f, "regulatory_json"))

    # 2. Internal raw files
    internal_files = list_adls_files("bronze/internal/", extensions=SUPPORTED_EXTENSIONS)
    for f in internal_files:
        tagged_inputs.append((f, "internal_raw"))

    # 3. Any other bronze files (auto-detect)
    other_files = list_adls_files("bronze/", extensions=SUPPORTED_EXTENSIONS)
    seen_paths = {t[0] for t in tagged_inputs}
    for f in other_files:
        if f not in seen_paths:
            tagged_inputs.append((f, "auto"))

    log_and_print(f"Discovered {len(tagged_inputs)} files to ingest")
    return tagged_inputs


def get_files_to_process(filepaths: list, filenames: list, file_bytes: list) -> list:
    """Filter files that need processing (not already indexed)."""
    files_to_process = []
    for filepath, filename, fb in zip(filepaths, filenames, file_bytes):
        files_to_process.append((filepath, filename, fb))
    return files_to_process

# COMMAND ----------
# MAGIC %md
# MAGIC ## Token Manager (Rate Limiting)

# COMMAND ----------
class TokenManager:
    """Simple token/rate manager for API calls."""
    def __init__(self, max_tokens_per_minute=80000):
        self.max_tokens = max_tokens_per_minute
        self.tokens_used = 0
        self.last_reset = time.time()

    def check_and_wait(self, estimated_tokens=1000):
        """Check if we can proceed, wait if rate limited."""
        now = time.time()
        if now - self.last_reset > 60:
            self.tokens_used = 0
            self.last_reset = now

        if self.tokens_used + estimated_tokens > self.max_tokens:
            wait_time = 60 - (now - self.last_reset)
            if wait_time > 0:
                log_and_print(f"  Rate limit approaching. Waiting {wait_time:.0f}s...")
                time.sleep(wait_time)
            self.tokens_used = 0
            self.last_reset = time.time()

        self.tokens_used += estimated_tokens

# COMMAND ----------
# MAGIC %md
# MAGIC ## Main Processing Pipeline — Load and Split Documents

# COMMAND ----------
def load_and_split_documents(files_to_process, token_manager, extracted_metadata=None):
    """Load and split documents from files based on their type.

    Routes each file to the appropriate parser based on extension.
    Returns a list of doc_records ready for chunking.
    """
    extracted_metadata = extracted_metadata or {}
    all_doc_records = []

    for filepath, filename, file_content in files_to_process:
        ext = os.path.splitext(filename)[1].lower()
        log_and_print(f"  Loading and splitting: {filename} ({ext})")

        try:
            # Save to temporary file for processing
            temp_file = None
            if isinstance(file_content, bytes):
                temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=ext)
                temp_file.write(file_content)
                temp_file.close()
                temp_path = temp_file.name
            else:
                temp_path = filepath

            # Route to appropriate processor
            if ext == ".pdf":
                records = process_pdf_with_ocr(
                    temp_path, file_content, filepath, token_manager,
                    extracted_metadata=extracted_metadata,
                )
            elif ext in (".docx", ".doc"):
                records = process_docx_with_metadata(
                    temp_path, file_content, filepath, token_manager,
                    extracted_metadata=extracted_metadata,
                )
            elif ext in (".html", ".htm"):
                if isinstance(file_content, bytes):
                    file_content_str = file_content.decode("utf-8", errors="replace")
                else:
                    file_content_str = file_content
                records = process_html_with_ocr(
                    temp_path, file_content_str, filepath, token_manager,
                    extracted_metadata=extracted_metadata,
                )
            elif ext in (".xlsx", ".xls", ".csv", ".tsv"):
                records = process_excel_with_metadata(
                    temp_path, file_content, filepath, token_manager,
                    extracted_metadata=extracted_metadata,
                )
            elif ext == ".json":
                records = process_json_document(
                    temp_path, file_content, filepath, token_manager,
                    extracted_metadata=extracted_metadata,
                )
            elif ext in (".txt", ".md"):
                records = process_text_file(
                    temp_path, file_content, filepath, token_manager,
                    extracted_metadata=extracted_metadata,
                )
            elif ext in (".png", ".jpg", ".jpeg", ".tiff", ".tif", ".bmp", ".gif"):
                records = process_image_file(
                    temp_path, file_content, filepath, token_manager,
                    extracted_metadata=extracted_metadata,
                )
            else:
                log_and_print(f"    Unsupported file type: {ext}", "warning")
                records = []

            all_doc_records.extend(records)

            # Clean up temp file
            if temp_file and os.path.exists(temp_path):
                os.unlink(temp_path)

        except Exception as e:
            log_and_print(f"    Error processing {filename}: {e}", "error")
            config.files_errored += 1
            continue

    return all_doc_records


def process_one_path(filepath: str, path_category: str = "auto"):
    """Process a single file path through the full pipeline.

    Steps:
    1. Read file from ADLS
    2. Route to appropriate parser
    3. Chunk the content
    4. Build semantic headers (Rule 1)
    5. Embed and index chunks (Rule 2)

    Returns summary dict with processing results.
    """
    token_manager = TokenManager()
    filename = os.path.basename(filepath)
    ext = os.path.splitext(filename)[1].lower()

    log_and_print(f"Processing [{path_category}]: {filepath}")

    # Step 1: Read from ADLS
    try:
        file_content = read_adls_bytes(filepath)
    except Exception as e:
        return {"path": filepath, "error": f"ADLS read failed: {e}"}

    # Step 2: Extract metadata if JSON metadata file exists alongside
    extracted_metadata = {}
    if path_category == "regulatory_json":
        # This IS the metadata JSON — parse it as regulatory content
        records = process_regulatory_scraped_json(
            filepath, file_content, filepath, token_manager,
        )
        if records:
            extracted_metadata = records[0].get("metadata", {})
    else:
        # Check for a companion metadata JSON
        metadata_path = filepath.rsplit(".", 1)[0] + "_metadata.json"
        try:
            meta_bytes = read_adls_bytes(metadata_path)
            meta_json = json.loads(meta_bytes)
            extracted_metadata = extract_metadata_from_json_content_enhanced(meta_json)
        except Exception:
            pass  # No companion metadata file

    # Step 3: Parse the file
    if path_category == "regulatory_json":
        doc_records = process_regulatory_scraped_json(
            filepath, file_content, filepath, token_manager,
            extracted_metadata=extracted_metadata,
        )
    else:
        files_to_process = [(filepath, filename, file_content)]
        doc_records = load_and_split_documents(
            files_to_process, token_manager,
            extracted_metadata=extracted_metadata,
        )

    if not doc_records:
        return {"path": filepath, "error": "No content extracted"}

    # Step 4: Chunk each document record
    all_chunks = []
    for doc_record in doc_records:
        content = doc_record.get("content", "")
        metadata = doc_record.get("metadata", {})
        # Merge doc-level fields into metadata
        for key in ("source_path", "source_file", "file_type", "extraction_method",
                     "has_tables", "has_images", "quality_score"):
            if key in doc_record:
                metadata[key] = doc_record[key]

        chunks = process_chunks(content, source_file=filename, metadata=metadata)
        all_chunks.extend(chunks)

    if not all_chunks:
        return {"path": filepath, "error": "No chunks produced"}

    # Step 5: Index chunks
    try:
        index_chunks_and_document({"metadata": extracted_metadata}, all_chunks)
    except Exception as e:
        return {"path": filepath, "error": f"Indexing failed: {e}", "chunk_count": len(all_chunks)}

    config.files_processed += 1
    config.total_chunks += len(all_chunks)

    return {
        "path": filepath,
        "path_category": path_category,
        "chunk_count": len(all_chunks),
        "doc_id": extracted_metadata.get("doc_id", filename),
        "regulator": extracted_metadata.get("regulator", ""),
        "document_class": extracted_metadata.get("document_class", ""),
    }

# COMMAND ----------
# MAGIC %md
# MAGIC ## Initialization — Set Up Stores

# COMMAND ----------
log_and_print("Step 1: Setting up Elasticsearch index and embedding model...")

try:
    es_client = get_es_client()
    create_es_index_with_mapping(es_client, config.es_index_name)
    log_and_print(f"Elasticsearch index '{config.es_index_name}' ready")
except Exception as e:
    log_and_print(f"ES setup failed: {e}", "error")
    es_client = None

try:
    create_azure_search_index(config)
except Exception as e:
    log_and_print(f"Azure Search setup failed (non-fatal): {e}", "warning")

log_and_print("Setup complete.")

# COMMAND ----------
# MAGIC %md
# MAGIC ## Execute Ingestion

# COMMAND ----------
# Discover files
input_paths_json = _widget("input_paths_json", "")
if not input_paths_json:
    tagged_inputs = discover_paths_to_ingest(config)
else:
    explicit_paths = json.loads(input_paths_json)
    tagged_inputs = [(p, "auto") for p in explicit_paths]
    log_and_print(f"Explicit paths: {len(tagged_inputs)}")

log_and_print(f"Total files to ingest: {len(tagged_inputs)}")
for sample_path, sample_cat in tagged_inputs[:10]:
    log_and_print(f"  [{sample_cat.upper()}] {sample_path}")
if len(tagged_inputs) > 10:
    log_and_print(f"  ... and {len(tagged_inputs) - 10} more")

# COMMAND ----------
# Process all files
results = []
for i, (path, category) in enumerate(tagged_inputs, start=1):
    log_and_print(f"\n[{i}/{len(tagged_inputs)}] [{category}] Processing: {path}")
    try:
        results.append(process_one_path(path, path_category=category))
    except Exception as exc:
        log_and_print(f"  ERROR: {exc}", "error")
        results.append({"path": path, "error": str(exc)})

# COMMAND ----------
# MAGIC %md
# MAGIC ## Ingestion Summary

# COMMAND ----------
success_count = sum(1 for r in results if "error" not in r)
error_count = sum(1 for r in results if "error" in r)
total_chunks = sum(r.get("chunk_count", 0) for r in results)

# Counts by path_category
category_counts = {}
regulator_counts = {}
class_counts = {}
for r in results:
    if "error" not in r:
        cat = r.get("path_category", "unknown")
        category_counts[cat] = category_counts.get(cat, 0) + 1
        reg = r.get("regulator") or "none"
        regulator_counts[reg] = regulator_counts.get(reg, 0) + 1
        dc = r.get("document_class", "unknown")
        class_counts[dc] = class_counts.get(dc, 0) + 1

print(f"""
{'='*60}
  NOVA INGESTION COMPLETE
{'='*60}
  Total processed:   {len(results)}
  Successful:        {success_count}
  Errors:            {error_count}
  {'─'*56}
  By path category:  {category_counts}
  By regulator:      {regulator_counts}
  By document_class: {class_counts}
  {'─'*56}
  Total chunks:      {total_chunks}
  Avg chunks/doc:    {total_chunks / max(success_count, 1):.1f}
{'='*60}
""")

if error_count:
    print("ERRORS:")
    for r in results:
        if "error" in r:
            print(f"  {r.get('path', 'unknown')}: {r['error']}")

# COMMAND ----------
# ---- Exit payload for orchestrator ----
exit_payload = json.dumps({
    "status": "success" if error_count == 0 else "partial_success",
    "total_processed": len(results),
    "successful": success_count,
    "errors": error_count,
    "total_chunks": total_chunks,
    "by_category": category_counts,
    "by_regulator": regulator_counts,
})
try:
    dbutils.notebook.exit(exit_payload)
except Exception:
    print(f"\nResult: {exit_payload}")
